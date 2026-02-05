#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
llm_only_extract_full.py

LLM-only baseline: 直接把整篇文档文本喂给 LLM，让模型自行抽取“学生疑问/困惑点”（不要求规范化）。
输出格式尽量对齐 question_candidates_v2_3.json，便于后续统一评测对比。

用法：
  python llm_only_extract_full.py <path_to_file_or_dir> --out llm_only_candidates.json

环境变量（任选其一配置）：
  - OPENAI_API_KEY        API Key
  - OPENAI_BASE_URL       兼容 OpenAI 的 base_url（例如 https://api.deepseek.com）
  - OPENAI_MODEL          模型名（默认：deepseek-chat 或你自己的兼容模型名）
  - OPENAI_TIMEOUT        请求超时秒数（默认 120）

说明：
  - 支持 .docx / .pdf
  - 不做“疑问段硬门控”：整篇输入，由 LLM 自行判断哪些是疑问/困惑点。
  - 为了兼容你现有流水线，本脚本输出 candidates 仍包含：
      source_line_no / sent_no / text / triggers / in_q_section / context_prev / context_next
    其中 source_line_no 尽量通过“文本匹配”回填；匹配不到则为 -1。
"""

import os
import re
import json
import argparse
from dataclasses import dataclass, asdict
from typing import List, Optional, Dict, Any, Tuple

OPENAI_API_KEY="sk-66df7e992c384ffd8d352fb948465d2e"
OPENAI_BASE_URL="https://api.deepseek.com"
OPENAI_MODEL="deepseek-chat"
OPENAI_TIMEOUT=120
# =====================
# 1) 文本读取：docx（含表格）/ pdf
# =====================

def read_docx_lines(path: str, max_paras: int = 1800, max_table_cells: int = 6000) -> List[str]:
    from docx import Document
    doc = Document(path)
    lines: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if len(lines) >= max_paras:
            break

    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = " ".join(
                    (pp.text or "").strip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()
                if cell_text:
                    row_parts.append(cell_text)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    return _clean_lines(lines)


def read_pdf_lines(path: str, max_pages: int = 12) -> List[str]:
    import fitz
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            if ln:
                lines.append(ln)
    return _clean_lines(lines)[:6000]


def _clean_lines(lines: List[str]) -> List[str]:
    cleaned = []
    for ln in lines:
        ln = (ln or "").replace("\u3000", " ")
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned


def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")


def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")) and not n.startswith("~$"):
                files.append(os.path.join(root, n))
    return sorted(files)

# =====================
# 2) LLM 调用（OpenAI 兼容）
# =====================

def _get_env(name: str, default: Optional[str] = None) -> Optional[str]:
    v = os.getenv(name)
    return v if (v is not None and str(v).strip() != "") else default


def call_llm_extract(full_text: str, file_name: str) -> List[Dict[str, Any]]:
    """
    让 LLM 返回一个 JSON 数组，每项至少包含：
      - question: str   （抽取到的疑问/困惑点原文表达，可是陈述句也可）
      - rationale: str  （为什么认为这是疑问/困惑）
      - evidence: str   （从文档中截取的原文证据片段，尽量短）
      - confidence: float in [0,1]
    """
    api_key = OPENAI_API_KEY
    base_url = OPENAI_BASE_URL
    model = OPENAI_MODEL
    timeout = OPENAI_TIMEOUT

    if not api_key:
        raise RuntimeError("Missing OPENAI_API_KEY in environment variables.")

    # 延迟导入，避免你不用时依赖报错
    import requests

    # OpenAI-compatible chat completions
    url = base_url.rstrip("/") + "/v1/chat/completions"

    system = (
        "你是“程序设计基础（C语言）”课程的助教。"
        "你的任务是从给定的学生自学报告全文中，抽取所有“疑问/困惑点”。\n"
        "注意：\n"
        "1) 不要总结心得，不要抽取纯陈述/学习收获。\n"
        "2) 既要抽取明确问句（含问号/疑问词），也要抽取“我不懂/不理解/困惑/不会/不知道”等困惑表达。\n"
        "3) 每个疑问尽量短，保留原意；不要进行规范化、不要补充新知识。\n"
        "4) 输出必须是严格 JSON（不要 Markdown，不要解释文字），格式为数组。\n"
        "5) 如果全文确实没有疑问，输出空数组 []。\n"
    )

    user = f"""文件名：{file_name}

【文档全文（按行拼接）】
{full_text}

【输出JSON数组格式示例】
[
  {{
    "question": "不清楚头文件以及部分函数和符号的用法及含义",
    "rationale": "出现了“不清楚…用法及含义”，表达困惑点",
    "evidence": "不清楚头文件以及部分函数和符号的用法及含义",
    "confidence": 0.78
  }}
]
"""

    payload = {
        "model": model,
        "temperature": 0.0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    }

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if r.status_code >= 400:
        raise RuntimeError(f"LLM HTTP {r.status_code}: {r.text[:800]}")

    data = r.json()
    content = data["choices"][0]["message"]["content"]

    # 尝试解析 JSON；若模型意外带了前后文字，做一次截断修复
    extracted = _safe_json_loads(content)
    if not isinstance(extracted, list):
        raise ValueError("LLM output is not a JSON array.")
    return extracted


def _safe_json_loads(s: str):
    s = (s or "").strip()
    try:
        return json.loads(s)
    except Exception:
        # 尝试抓取第一个 '[' 到最后一个 ']'
        a = s.find("[")
        b = s.rfind("]")
        if a != -1 and b != -1 and b > a:
            return json.loads(s[a:b+1])
        raise

# =====================
# 3) 回填到候选结构（对齐 v2.3 candidates）
# =====================

@dataclass
class Candidate:
    source_line_no: int
    sent_no: int
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]
    llm: Optional[Dict[str, Any]] = None


def _locate_in_lines(lines: List[str], evidence: str) -> int:
    """
    在 lines 中用“弱匹配”定位 evidence 所在行号（0-based）。
    找不到返回 -1。
    """
    if not evidence:
        return -1
    ev = re.sub(r"\s+", " ", evidence).strip()
    if len(ev) < 2:
        return -1

    # 先精确包含匹配
    for i, ln in enumerate(lines):
        if ev in ln:
            return i

    # 再用简化匹配：取前 N 字
    head = ev[:12]
    if len(head) >= 4:
        for i, ln in enumerate(lines):
            if head in ln:
                return i

    return -1


def build_candidates(lines: List[str], llm_items: List[Dict[str, Any]]) -> List[Candidate]:
    cands: List[Candidate] = []

    for idx, it in enumerate(llm_items):
        q = (it.get("question") or "").strip()
        if not q:
            continue

        evidence = (it.get("evidence") or q).strip()
        line_no = _locate_in_lines(lines, evidence)

        prev_line = lines[line_no - 1].strip() if (line_no is not None and line_no >= 1) else None
        next_line = lines[line_no + 1].strip() if (line_no is not None and line_no >= 0 and line_no + 1 < len(lines)) else None

        # triggers 仅用于兼容；后续评测可用 triggers 区分来源
        triggers = ["LLM_ONLY"]
        # 这里不做 Q-section 判断，统一设 False（因为是“整篇输入”的基线）
        cands.append(Candidate(
            source_line_no=line_no if line_no is not None else -1,
            sent_no=0,
            text=q.replace("?", "？"),
            triggers=triggers,
            in_q_section=False,
            context_prev=prev_line,
            context_next=next_line,
            llm={
                "rationale": it.get("rationale", ""),
                "evidence": evidence,
                "confidence": it.get("confidence", None),
            }
        ))

    # 去重（按 text）
    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        key = c.text
        if key not in seen:
            uniq.append(c)
            seen.add(key)

    return uniq

# =====================
# 4) main
# =====================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path", default="norm_dataset",help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="llm_only_candidates.json", help="output json filename")
    parser.add_argument("--max_chars", type=int, default=24000, help="truncate concatenated text to avoid overlong prompts")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file")
    args = parser.parse_args()

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)

            # 拼接全文给 LLM（避免过长：截断）
            joined = "\n".join(lines)
            if len(joined) > args.max_chars:
                joined = joined[:args.max_chars] + "\n(后续内容因长度限制已截断)"

            llm_items = call_llm_extract(joined, os.path.basename(p))
            cands = build_candidates(lines, llm_items)

            item = {
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(cands),
                "candidates": [asdict(x) for x in cands],
            }
            results.append(item)

            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(cands)}")

            if args.max_preview > 0:
                for x in cands[:args.max_preview]:
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) {x.text}  triggers={x.triggers} conf={x.llm.get('confidence') if x.llm else None}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {args.out} ({len(results)} files)")


if __name__ == "__main__":
    main()
