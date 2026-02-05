#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
question_extract_v2_3_fallback_llm.py

规则优先 + 0候选时 LLM strict 兜底。
新增：支持“问题/疑问/困惑/疑惑”作为单独标题行后，抽取其后续枚举条目。
修复：marker_block 支持同一行出现多个编号条目（如 "1.xxx 2.yyy"）。
"""

import os
import re
import json
import argparse
from dataclasses import dataclass, asdict
from typing import List, Optional, Tuple, Dict, Any

OPENAI_API_KEY="sk-66df7e992c384ffd8d352fb948465d2e"
OPENAI_BASE_URL="https://api.deepseek.com"
OPENAI_MODEL="deepseek-chat"
OPENAI_TIMEOUT=120

# =====================
# 0) 工具：输出避免重名
# =====================

def make_nonconflicting_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if not ext:
        ext = ".json"
    return f"{base}_{ts}{ext}"

# =====================
# 1) 文本读取：docx/pdf
# =====================

def _clean_lines(lines: List[str]) -> List[str]:
    cleaned = []
    for ln in lines:
        ln = (ln or "").replace("\u3000", " ")
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned

def read_docx_lines(path: str, max_paras: int = 1200, max_table_cells: int = 4000) -> List[str]:
    from docx import Document
    doc = Document(path)

    lines: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if len(lines) >= max_paras:
            break

    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = " ".join(
                    (pp.text or "").strip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()
                if cell_text:
                    row_parts.append(cell_text)
                cell_count += 1
                if cell_count >= max_table_cells:
                    break
            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)
            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    return _clean_lines(lines)

def read_pdf_lines(path: str, max_pages: int = 8) -> List[str]:
    import fitz
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            ln = ln.replace("\u3000", " ")
            ln = re.sub(r"\s+", " ", ln).strip()
            if ln:
                lines.append(ln)
    return lines[:4000]

def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")

def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")) and not n.startswith("~$"):
                files.append(os.path.join(root, n))
    return sorted(files)

# =====================
# 2) v2.3 规则抽取
# =====================

QUESTION_WORDS = [
    "为什么", "为何", "怎么", "如何", "怎样", "是否", "能否", "可否",
    "区别", "联系", "原因", "作用", "意义", "什么时候", "哪里", "多少",
    "是什么", "什么", "吗"
]

CONFUSION_WORDS = [
    "不懂", "没懂", "不理解", "没理解", "不太懂", "不太明白", "不清楚",
    "有点懵", "困惑", "疑惑", "看不懂", "搞不懂", "不确定", "不熟悉",
    "不会", "不知道", "不明白", "搞混", "不了解"
]

MARKER_WORDS = [
    "疑问", "问题", "困惑", "疑惑", "不懂", "不会", "没懂", "不理解", "不明白", "不清楚",
    "遇到的问题", "存在的问题", "有困难"
]

Q_SECTION_HINTS = [
    "自学心得或疑问", "心得或疑问", "心得与疑问", "反思与问题",
    "疑问", "问题", "困惑", "疑惑",
]
EXCLUDE_SECTION_HINTS = [
    "自学内容", "学习内容", "自学过程", "学习过程", "自学总结", "学习总结",
]
END_SECTION_HINTS = [
    "自学照片", "照片", "截图", "图片", "附件", "附录",
]

NEGATIVE_PATTERNS = [
    r"^目录$",
    r"^contents$",
    r"^第[一二三四五六七八九十0-9]{1,2}[章节]\b",
    r"^[0-9]+\s*$",
]

META_FIELD_HINTS = ["组长", "组员", "姓名", "学号", "班级", "日期", "学院", "专业"]

WEAK_QUESTION_HINTS = ["含义", "意思", "代表", "表示", "用途", "用法", "作用"]
WEAK_CONFUSE_PAT = re.compile(r"(含义|意思|用法|用途|作用).{0,6}(不明确|不清楚|不太明白|不明白|不理解|不了解)")

COURSE_SIGNAL_WORDS = [
    "for", "while", "do", "if", "else", "switch", "case", "break", "continue",
    "scanf", "printf", "gets", "puts", "fopen", "fclose", "fprintf", "fscanf",
    "%d", "%f", "%c", "%s", "\\t", "\\n", "\\0",
    "数组", "字符串", "字符数组", "指针", "函数", "递归", "循环", "分支", "进制",
    "变量", "常量", "数据类型", "结构体", "链表", "文件", "输入", "输出",
    "ASCII", "float", "double", "int", "char"
]

# “问题/疑问/困惑/疑惑”标题行：允许带不带冒号
MARKER_HEADER_RE = re.compile(r"^\s*(疑问|问题|困惑|疑惑)\s*[:：]?\s*$")

ENUM_MARK_RE = re.compile(r"(?:^|\s)(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*")

@dataclass
class Candidate:
    source_line_no: int
    sent_no: int
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]
    llm: Optional[Dict[str, Any]] = None

def is_negative_line(line: str) -> bool:
    l = line.strip().lower()
    for pat in NEGATIVE_PATTERNS:
        if re.search(pat, l, flags=re.IGNORECASE):
            return True
    return False

def looks_like_section_title(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if "？" in s or any(w in s for w in ["如何", "怎么", "为什么", "是什么", "是否", "能否", "可否", "吗"]):
        return False

    title_like = s.replace("：", "").replace(":", "").strip()
    if title_like in Q_SECTION_HINTS or title_like in EXCLUDE_SECTION_HINTS or title_like in END_SECTION_HINTS:
        return True

    if re.match(r"^\s*(第?\s*[一二三四五六七八九十0-9]{1,2}\s*[、\.．:：\)])\s*", s):
        return True
    if re.match(r"^\s*(\(?\d{1,2}\)?\s*[、\.．:：\)])\s*", s):
        return True
    return False

def update_section_flag(line: str, in_q_section: bool) -> bool:
    s = line.strip()
    if any(h in s for h in END_SECTION_HINTS):
        return False
    if any(h in s for h in EXCLUDE_SECTION_HINTS):
        return False
    if any(h in s for h in Q_SECTION_HINTS):
        return True
    return in_q_section

def split_into_sentences_keep_q(line: str) -> List[str]:
    if not line:
        return []
    s = re.sub(r"\s+", " ", line).strip()
    if not s:
        return []
    s = s.replace("?", "？")

    parts: List[str] = []
    start = 0
    for m in re.finditer(r"[。；;！!？]\s*", s):
        end = m.end()
        seg = s[start:end].strip()
        if seg:
            parts.append(seg)
        start = end
    last = s[start:].strip()
    if last:
        parts.append(last)

    final_parts: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        sub = re.split(r"(?:\s+|^)(?=(?:\(?\d{1,2}\)?[\.、\)])\s*)", p)
        sub = [x.strip() for x in sub if x.strip()]
        if len(sub) <= 1:
            final_parts.append(p)
        else:
            buf = ""
            for x in sub:
                if re.fullmatch(r"\(?\d{1,2}\)?[\.、\)]", x):
                    buf = x
                else:
                    if buf:
                        final_parts.append((buf + " " + x).strip())
                        buf = ""
                    else:
                        final_parts.append(x)
            if buf:
                final_parts.append(buf)
    return [x for x in final_parts if len(x) >= 2]

def split_marker_tail(tail: str) -> List[str]:
    tail = (tail or "").strip()
    if not tail:
        return []
    tail = tail.replace("?", "？")
    parts = split_into_sentences_keep_q(tail) or [tail]

    refined: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        idxs = [m.start() for m in ENUM_MARK_RE.finditer(p)]
        if len(idxs) <= 1:
            refined.append(p)
            continue
        idxs.append(len(p))
        for a, b in zip(idxs, idxs[1:]):
            seg = p[a:b].strip()
            seg = re.sub(r"^\s*(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*", "", seg)
            seg = seg.strip()
            if seg:
                refined.append(seg)

    out = []
    for x in refined:
        x = re.sub(r"\s+", " ", x).strip()
        if "？" not in x and any(w in x for w in ["什么", "怎么", "如何", "为什么", "是否", "能否", "可否"]):
            x = x + "？"
        if len(x) >= 2:
            out.append(x)
    return out

def extract_marker_contents(text: str) -> List[str]:
    t = re.sub(r"\s*\|\s*", " | ", text)
    m = re.search(
        rf"({'|'.join(map(re.escape, MARKER_WORDS))})\s*(?:[:：]\s*|\|\s*[:：]?\s*\|\s*)(.+)$",
        t
    )
    if not m:
        return []
    tail = m.group(2).strip()
    tail = re.sub(r"^[\|\s:：]+", "", tail)
    tail = re.sub(r"\s*\|\s*", " ", tail).strip()
    if not tail:
        return []
    return split_marker_tail(tail)

def score_triggers(sent: str, in_q_section: bool) -> Tuple[List[str], int]:
    triggers: List[str] = []
    score = 0
    sent2 = sent.replace("?", "？")

    if "？" in sent2:
        triggers.append("HAS_QUESTION_MARK")
        score += 3

    for w in QUESTION_WORDS:
        if w in sent2:
            triggers.append(f"QWORD:{w}")
            score += 2
            break

    for w in CONFUSION_WORDS:
        if w in sent2:
            triggers.append(f"CONFUSE:{w}")
            score += 2 if not in_q_section else 3
            break

    if any(h in sent2 for h in WEAK_QUESTION_HINTS):
        triggers.append("WEAK_HINT")
        score += 1 if in_q_section else 0

    if in_q_section and score > 0:
        score += 1
        triggers.append("IN_Q_SECTION")

    return triggers, score

def extract_candidates_v2_3(lines: List[str]) -> List[Candidate]:
    cands: List[Candidate] = []
    in_q_section = False
    marker_block = False

    for i, raw_line in enumerate(lines):
        line = (raw_line or "").strip()
        if not line:
            marker_block = False
            continue
        if is_negative_line(line):
            marker_block = False
            continue

        if looks_like_section_title(line):
            in_q_section = update_section_flag(line, in_q_section)
            marker_block = False
            continue

        if not in_q_section:
            marker_block = False
            continue

        prev_line = lines[i - 1].strip() if i - 1 >= 0 else None
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None

        # ✅ 单独标题行：问题/疑问/困惑/疑惑（可带冒号）
        if MARKER_HEADER_RE.match(line):
            marker_block = True
            continue

        # ✅ marker_block：支持同一行出现多个编号（关键修复）
        if marker_block:
            # 只要这一行出现编号，就拆开抽取
            if ENUM_MARK_RE.search(line):
                items = split_marker_tail(line)
                for j, it in enumerate(items):
                    it = (it or "").strip().replace("?", "？")
                    if len(it) < 3:
                        continue
                    triggers, _ = score_triggers(it, in_q_section=True)
                    triggers = sorted(list(set((triggers or []) + ["MARKER:IMPLICIT", "IN_Q_SECTION"])))
                    cands.append(Candidate(
                        source_line_no=i,
                        sent_no=j,
                        text=it,
                        triggers=triggers,
                        in_q_section=True,
                        context_prev=prev_line,
                        context_next=next_line,
                        llm=None
                    ))
                continue
            else:
                # 这一行没有任何编号 -> 退出列表模式，继续走正常规则
                marker_block = False

        if any(f in line for f in META_FIELD_HINTS):
            has_qmark = ("?" in line or "？" in line)
            has_marker = bool(extract_marker_contents(line))
            if not has_qmark and not has_marker:
                continue

        marker_parts = extract_marker_contents(line)
        if marker_parts:
            for j, p in enumerate(marker_parts):
                p = p.strip()
                if not p:
                    continue
                triggers, _ = score_triggers(p, in_q_section=True)
                triggers = sorted(list(set((triggers or []) + ["MARKER:EXPLICIT", "IN_Q_SECTION"])))
                cands.append(Candidate(
                    source_line_no=i,
                    sent_no=j,
                    text=p.replace("?", "？"),
                    triggers=triggers,
                    in_q_section=True,
                    context_prev=prev_line,
                    context_next=next_line,
                    llm=None
                ))
            marker_block = True
            continue

        sents = split_into_sentences_keep_q(line)
        for j, sent in enumerate(sents):
            sent = sent.strip().replace("?", "？")
            if not sent:
                continue

            if WEAK_CONFUSE_PAT.search(sent) and "？" not in sent and not any(w in sent for w in QUESTION_WORDS):
                continue

            if len(sent) <= 4 and ("？" not in sent):
                continue

            triggers, score = score_triggers(sent, in_q_section=True)
            if score < 2:
                continue

            only_weak = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                (any(t.startswith("CONFUSE:") for t in triggers) or "WEAK_HINT" in triggers)
            )

            confuse_only = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                any(t.startswith("CONFUSE:") for t in triggers)
            )

            sent_l = sent.lower()
            has_course_signal = any(k.lower() in sent_l for k in COURSE_SIGNAL_WORDS)

            if only_weak:
                if confuse_only and has_course_signal and len(sent) >= 8:
                    triggers.append("ALLOW:CONFUSE_WITH_SIGNAL")
                else:
                    continue

            cands.append(Candidate(
                source_line_no=i,
                sent_no=j,
                text=sent,
                triggers=sorted(list(set(triggers))),
                in_q_section=True,
                context_prev=prev_line,
                context_next=next_line,
                llm=None
            ))

    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        if c.text not in seen:
            uniq.append(c)
            seen.add(c.text)
    return uniq

# =====================
# 3) strict LLM 兜底抽取（允许输出 []）
# =====================

def _safe_json_loads(s: str):
    s = (s or "").strip()
    try:
        return json.loads(s)
    except Exception:
        a = s.find("[")
        b = s.rfind("]")
        if a != -1 and b != -1 and b > a:
            return json.loads(s[a:b + 1])
        raise

def call_llm_extract_strict(full_text: str, file_name: str, timeout: int = 120) -> List[Dict[str, Any]]:
    api_key = OPENAI_API_KEY
    base_url = OPENAI_BASE_URL
    model = OPENAI_MODEL

    if not api_key:
        raise RuntimeError("Missing OPENAI_API_KEY in environment variables.")

    import requests
    url = base_url.rstrip("/") + "/v1/chat/completions"

    system = (
        "你是信息抽取助手。任务：从学生自学报告全文中抽取“可用于RAG问答”的疑问/困惑点。\n"
        "【只允许抽取两类】\n"
        "A) QUESTION：明确问句，满足其一即可：\n"
        "   - 含问号（? 或 ？）\n"
        "   - 或包含疑问词：什么/怎么/如何/为什么/是否/能否/可否/区别/联系/原因/作用/意义/什么时候/哪里/多少/吗\n"
        "B) CONFUSION：明确困惑陈述，必须同时满足：\n"
        "   - 含困惑/不会表达：不懂/不理解/没理解/不太懂/不太明白/不清楚/疑惑/困惑/搞不懂/不知道/不确定/不熟悉/把握不好/掌握不好/卡壳\n"
        "   - 且句子中指向具体对象（概念/语法/函数/例子/知识点），不是空泛的“学得不好/有困难”。\n"
        "【严格排除】\n"
        "1) 只列出主题/目录/知识点列表但没有问法或困惑动词，一律不要抽取。\n"
        "2) 纯心得/收获/总结/计划，不要抽取。\n"
        "3) 过于笼统无法形成问答的句子默认不抽取。\n"
        "【输出要求】\n"
        "- 只输出严格JSON数组。\n"
        "- evidence 必须是原文片段（<=60字优先）。\n"
        "- question 尽量贴近原文，不要规范化。\n"
        "- 若没有符合条件的内容，输出 []。\n"
    )

    user = f"""文件名：{file_name}

【文档全文（按行拼接）】
{full_text}

【输出JSON数组格式示例】
[
  {{
    "question": "对二进制、十进制、八进制、十六进制的换算不懂",
    "evidence": "对二进制、十进制、八进制、十六进制的换算不懂",
    "category": "CONFUSION",
    "confidence": 0.9,
    "rationale": "包含“不懂”且指向具体知识点“进制换算”"
  }}
]
"""

    payload = {
        "model": model,
        "temperature": 0.0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    }

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if r.status_code >= 400:
        raise RuntimeError(f"LLM HTTP {r.status_code}: {r.text[:800]}")

    data = r.json()
    content = data["choices"][0]["message"]["content"]
    extracted = _safe_json_loads(content)
    if not isinstance(extracted, list):
        raise ValueError("LLM output is not a JSON array.")
    return extracted

def _locate_in_lines(lines: List[str], evidence: str) -> int:
    if not evidence:
        return -1
    ev = re.sub(r"\s+", " ", evidence).strip()
    if len(ev) < 2:
        return -1
    for i, ln in enumerate(lines):
        if ev in ln:
            return i
    head = ev[:12]
    if len(head) >= 4:
        for i, ln in enumerate(lines):
            if head in ln:
                return i
    return -1

def build_llm_candidates(lines: List[str], llm_items: List[Dict[str, Any]]) -> List[Candidate]:
    cands: List[Candidate] = []
    for it in llm_items:
        q = (it.get("question") or "").strip()
        if not q:
            continue
        evidence = (it.get("evidence") or q).strip()
        line_no = _locate_in_lines(lines, evidence)

        prev_line = lines[line_no - 1].strip() if (line_no >= 1) else None
        next_line = lines[line_no + 1].strip() if (line_no >= 0 and line_no + 1 < len(lines)) else None

        cands.append(Candidate(
            source_line_no=line_no,
            sent_no=0,
            text=q.replace("?", "？"),
            triggers=["LLM_FALLBACK_STRICT"],
            in_q_section=False,
            context_prev=prev_line,
            context_next=next_line,
            llm={
                "category": it.get("category", ""),
                "rationale": it.get("rationale", ""),
                "evidence": evidence,
                "confidence": it.get("confidence", None),
            }
        ))

    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        if c.text not in seen:
            uniq.append(c)
            seen.add(c.text)
    return uniq

# =====================
# 4) 主流程
# =====================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="question_candidates_v2_3_fallback.json", help="output json filename")
    parser.add_argument("--max_chars", type=int, default=24000, help="truncate text for LLM prompt")
    parser.add_argument("--llm_timeout", type=int, default=OPENAI_TIMEOUT, help="LLM request timeout seconds")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file for quick check")
    args = parser.parse_args()

    out_path = make_nonconflicting_path(args.out)

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)

            rule_cands = extract_candidates_v2_3(lines)
            final_cands = rule_cands
            used_fallback = False

            if len(rule_cands) == 0:
                joined = "\n".join(lines)
                if len(joined) > args.max_chars:
                    joined = joined[:args.max_chars] + "\n(后续内容因长度限制已截断)"
                llm_items = call_llm_extract_strict(joined, os.path.basename(p), timeout=args.llm_timeout)
                llm_cands = build_llm_candidates(lines, llm_items)
                final_cands = llm_cands
                used_fallback = True

            results.append({
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(final_cands),
                "candidates": [asdict(x) for x in final_cands],
                "fallback_used": used_fallback,
            })

            tag = "RULE" if not used_fallback else "FALLBACK_LLM"
            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(final_cands)} mode={tag}")

            if args.max_preview > 0:
                for x in final_cands[:args.max_preview]:
                    extra = ""
                    if x.llm:
                        extra = f"  [cat={x.llm.get('category')}, conf={x.llm.get('confidence')}]"
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) {x.text} triggers={x.triggers}{extra}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {out_path} ({len(results)} files)")


if __name__ == "__main__":
    main()
