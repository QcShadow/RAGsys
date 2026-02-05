#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
llm_only_extract_full_strict.py

LLM-only baseline（改进提示词版）：
- 直接输入整篇文档文本给 LLM
- 只抽取“可用于RAG问答”的疑问/困惑点（更严格，避免把主题列表当问题）
- 不要求规范化
- 输出结构对齐 question_candidates_v2_3.json：file/num_lines/num_candidates/candidates

并且：输出文件自动避免重名（若已存在则追加时间戳）。
"""

import os
import re
import json
import argparse
from dataclasses import dataclass, asdict
from typing import List, Optional, Dict, Any

OPENAI_API_KEY="sk-66df7e992c384ffd8d352fb948465d2e"
OPENAI_BASE_URL="https://api.deepseek.com"
OPENAI_MODEL="deepseek-chat"
OPENAI_TIMEOUT=120
# =====================
# 1) 读取 docx/pdf
# =====================

def _clean_lines(lines: List[str]) -> List[str]:
    cleaned = []
    for ln in lines:
        ln = (ln or "").replace("\u3000", " ")
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned


def read_docx_lines(path: str, max_paras: int = 2000, max_table_cells: int = 8000) -> List[str]:
    from docx import Document
    doc = Document(path)
    lines: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if len(lines) >= max_paras:
            break

    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = " ".join(
                    (pp.text or "").strip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()
                if cell_text:
                    row_parts.append(cell_text)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    return _clean_lines(lines)


def read_pdf_lines(path: str, max_pages: int = 12) -> List[str]:
    import fitz
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            if ln:
                lines.append(ln)
    return _clean_lines(lines)[:8000]


def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")


def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")) and not n.startswith("~$"):
                files.append(os.path.join(root, n))
    return sorted(files)


# =====================
# 2) LLM 调用（OpenAI-compatible）
# =====================

def _get_env(name: str, default: Optional[str] = None) -> Optional[str]:
    v = os.getenv(name)
    return v if (v is not None and str(v).strip() != "") else default


def _safe_json_loads(s: str):
    s = (s or "").strip()
    try:
        return json.loads(s)
    except Exception:
        a = s.find("[")
        b = s.rfind("]")
        if a != -1 and b != -1 and b > a:
            return json.loads(s[a:b+1])
        raise


def call_llm_extract(full_text: str, file_name: str) -> List[Dict[str, Any]]:
    """
    返回 JSON 数组，每项包含：
      - question: str        抽取到的疑问/困惑点（尽量贴近原文）
      - evidence: str        文档中的原文证据（短片段，用于定位/审计）
      - category: str        QUESTION / CONFUSION
      - confidence: float    [0,1]
      - rationale: str       简短理由
    """
    api_key = OPENAI_API_KEY
    base_url = OPENAI_BASE_URL
    model = OPENAI_MODEL
    timeout = OPENAI_TIMEOUT

    if not api_key:
        raise RuntimeError("Missing OPENAI_API_KEY in environment variables.")

    import requests
    url = base_url.rstrip("/") + "/v1/chat/completions"

    # ===== Prompt V2（严格版：避免主题列表误报）=====
    system = (
        "你是信息抽取助手。任务：从学生自学报告全文中抽取“可用于RAG问答”的疑问/困惑点。\n"
        "【只允许抽取两类】\n"
        "A) QUESTION：明确问句，满足其一即可：\n"
        "   - 含问号（? 或 ？）\n"
        "   - 或包含疑问词：什么/怎么/如何/为什么/是否/能否/可否/区别/联系/原因/作用/意义/什么时候/哪里/多少/吗\n"
        "B) CONFUSION：明确困惑陈述，必须同时满足：\n"
        "   - 含困惑/不会表达：不懂/不理解/没理解/不太懂/不太明白/不清楚/疑惑/困惑/搞不懂/不知道/不确定/不熟悉/把握不好/掌握不好/卡壳\n"
        "   - 且句子中指向具体对象（概念/语法/函数/例子/知识点），不是空泛的“学得不好/有困难”。\n"
        "【严格排除】\n"
        "1) 只列出主题/目录/知识点列表（如“变量和常量、运算符和表达式”）但没有问法或困惑动词，一律不要抽取。\n"
        "2) 纯心得/收获/总结/表扬/计划，不要抽取。\n"
        "3) 过于笼统无法形成问答的句子（如“编写代码有困难”）默认不抽取，除非后面紧跟具体困难点。\n"
        "【输出要求】\n"
        "- 只输出严格JSON数组（不要Markdown/不要解释）。\n"
        "- evidence 必须是从原文复制的一小段（<=60字优先），用于定位。\n"
        "- question 尽量贴近原文表达，不要做规范化和扩写。\n"
        "- 若没有符合条件的内容，输出 []。\n"
    )

    user = f"""文件名：{file_name}

【文档全文（按行拼接）】
{full_text}

【输出JSON数组格式示例】
[
  {{
    "question": "对二进制、十进制、八进制、十六进制的换算不懂",
    "evidence": "对二进制、十进制、八进制、十六进制的换算不懂",
    "category": "CONFUSION",
    "confidence": 0.9,
    "rationale": "包含“不懂”且指向具体知识点“进制换算”"
  }},
  {{
    "question": "printf括号中的“%d”作用是什么？",
    "evidence": "printf括号中的“%d”作用是什么",
    "category": "QUESTION",
    "confidence": 0.95,
    "rationale": "明确问句，包含“作用是什么”"
  }}
]
"""

    payload = {
        "model": model,
        "temperature": 0.0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    }

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if r.status_code >= 400:
        raise RuntimeError(f"LLM HTTP {r.status_code}: {r.text[:800]}")

    data = r.json()
    content = data["choices"][0]["message"]["content"]
    extracted = _safe_json_loads(content)
    if not isinstance(extracted, list):
        raise ValueError("LLM output is not a JSON array.")
    return extracted


# =====================
# 3) 对齐候选结构（兼容 v2.3）
# =====================

@dataclass
class Candidate:
    source_line_no: int
    sent_no: int
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]
    llm: Optional[Dict[str, Any]] = None


def _locate_in_lines(lines: List[str], evidence: str) -> int:
    if not evidence:
        return -1
    ev = re.sub(r"\s+", " ", evidence).strip()
    if len(ev) < 2:
        return -1

    for i, ln in enumerate(lines):
        if ev in ln:
            return i

    head = ev[:12]
    if len(head) >= 4:
        for i, ln in enumerate(lines):
            if head in ln:
                return i

    return -1


def build_candidates(lines: List[str], llm_items: List[Dict[str, Any]]) -> List[Candidate]:
    cands: List[Candidate] = []
    for it in llm_items:
        q = (it.get("question") or "").strip()
        if not q:
            continue
        evidence = (it.get("evidence") or q).strip()
        line_no = _locate_in_lines(lines, evidence)

        prev_line = lines[line_no - 1].strip() if (line_no >= 1) else None
        next_line = lines[line_no + 1].strip() if (line_no >= 0 and line_no + 1 < len(lines)) else None

        cands.append(Candidate(
            source_line_no=line_no,
            sent_no=0,
            text=q.replace("?", "？"),
            triggers=["LLM_ONLY_STRICT"],
            in_q_section=False,
            context_prev=prev_line,
            context_next=next_line,
            llm={
                "category": it.get("category", ""),
                "rationale": it.get("rationale", ""),
                "evidence": evidence,
                "confidence": it.get("confidence", None),
            }
        ))

    # 去重（按 text）
    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        if c.text not in seen:
            uniq.append(c)
            seen.add(c.text)
    return uniq


# =====================
# 4) 输出避免重名
# =====================

def make_nonconflicting_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if not ext:
        ext = ".json"
    return f"{base}_{ts}{ext}"


# =====================
# 5) main
# =====================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="llm_only_candidates_strict.json", help="output json filename")
    parser.add_argument("--max_chars", type=int, default=24000, help="truncate text to avoid overlong prompts")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file")
    args = parser.parse_args()

    out_path = make_nonconflicting_path(args.out)

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)
            joined = "\n".join(lines)
            if len(joined) > args.max_chars:
                joined = joined[:args.max_chars] + "\n(后续内容因长度限制已截断)"

            llm_items = call_llm_extract(joined, os.path.basename(p))
            cands = build_candidates(lines, llm_items)

            results.append({
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(cands),
                "candidates": [asdict(x) for x in cands],
            })

            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(cands)}")

            if args.max_preview > 0:
                for x in cands[:args.max_preview]:
                    conf = x.llm.get("confidence") if x.llm else None
                    cat = x.llm.get("category") if x.llm else None
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) [{cat}] {x.text}  conf={conf}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {out_path} ({len(results)} files)")


if __name__ == "__main__":
    main()
