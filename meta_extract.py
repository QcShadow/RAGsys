import os
import re
import json
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple

# -------- 可调：课程主题关键词（你后续可继续补充/细化）--------
TOPIC_KEYWORDS: Dict[str, List[str]] = {
    "概述/基础": ["程序设计基础", "C语言程序设计", "开发过程", "IDE", "关键字", "标识符"],
    "数据类型/输入输出": ["数据类型", "整型", "浮点", "double", "float", "char", "scanf", "printf", "输入输出", "格式控制"],
    "分支结构": ["if", "else", "switch", "分支", "选择结构", "条件判断"],
    "循环结构": ["for", "while", "do-while", "循环", "迭代", "break", "continue"],
    "数组": ["数组", "一维数组", "二维数组", "多维数组", "下标", "字符串", "字符数组"],
    "函数": ["函数", "return", "形参", "实参", "递归", "作用域", "存储类型"],
    "指针": ["指针", "地址", "&", "*", "指针变量", "指针数组", "函数指针"],
    "结构体/链表": ["结构体", "struct", "union", "enum", "typedef", "链表", "malloc", "free"],
    "文件": ["文件", "fopen", "fclose", "fprintf", "fscanf", "fread", "fwrite", "fgetc", "fputc"],
}

DOC_TYPE_RULES: List[Tuple[str, List[str]]] = [
    ("自学报告", ["自学报告", "自主学习", "自学内容", "自学心得"]),
    ("预习报告", ["预习报告", "预习内容", "课前预习"]),
    ("实验报告", ["实验报告", "实验内容", "实验步骤", "实验结果"]),
    ("作业/大作业", ["大作业", "课程设计", "作业要求", "作业内容"]),
]

COURSE_HINTS = ["程序设计基础", "C语言程序设计", "C 语言程序设计", "程序设计"]

CN_NUM = {"一":1,"二":2,"三":3,"四":4,"五":5,"六":6,"七":7,"八":8,"九":9,"十":10}

# ------------- 读取 docx / pdf 文本 -------------
def read_docx_text(path: str, max_paras: int = 200, max_table_cells: int = 400) -> List[str]:
    """
    读取 docx 的段落 + 表格文本。
    返回一组“行”，用于后续标题/元数据/主题抽取。
    """
    from docx import Document  # python-docx
    doc = Document(path)

    lines: List[str] = []

    # 1) 段落
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if len(lines) >= max_paras:
            break

    # 2) 表格（关键：组长/组员/周次常在这里）
    cell_count = 0
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_text_parts = []
            for ci, cell in enumerate(row.cells):
                cell_text = " ".join(
                    (pp.text or "").strip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()

                if cell_text:
                    row_text_parts.append(cell_text)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            # 把一行表格合成一行文本：用 | 分隔，利于规则匹配
            row_text = " | ".join(row_text_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    # 3) 去重（保留顺序）
    deduped = []
    seen = set()
    for ln in lines:
        if ln not in seen:
            deduped.append(ln)
            seen.add(ln)

    return deduped

def read_pdf_text(path: str, max_pages: int = 2) -> List[str]:
    # 用 PyMuPDF（fitz）读取前几页文本；一般环境都装得上
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        # 切成行，保留前面多一些
        for ln in text.splitlines():
            ln = ln.strip()
            if ln:
                lines.append(ln)
    return lines[:500]

# ------------- 抽取逻辑 -------------
def guess_doc_type(text: str) -> Tuple[str, List[str]]:
    hits = []
    for dt, kws in DOC_TYPE_RULES:
        for kw in kws:
            if kw in text:
                hits.append(kw)
        if hits:
            return dt, hits[:6]
    return "未知", []

def guess_course(text: str) -> Optional[str]:
    for h in COURSE_HINTS:
        if h in text:
            return h
    return None

def guess_week(text: str) -> Optional[int]:
    m = re.search(r"第\s*([0-9]{1,2})\s*周", text)
    if m:
        return int(m.group(1))

    m = re.search(r"第\s*([一二三四五六七八九十]{1,2})\s*周", text)
    if m:
        s = m.group(1)
        if s == "十":
            return 10
        if len(s) == 2 and s[0] == "十":   # 十一/十二…
            return 10 + CN_NUM.get(s[1], 0)
        if len(s) == 2 and s[1] == "十":   # 二十这种一般不会出现周次，但兼容
            return CN_NUM.get(s[0], 0) * 10
        return CN_NUM.get(s, None)

    m = re.search(r"\bWeek\s*([0-9]{1,2})\b", text, re.IGNORECASE)
    if m:
        return int(m.group(1))

    return None

def guess_title(lines: List[str]) -> Optional[str]:
    # 简单：取前 10 行里最长且“像标题”的那一行
    cand = []
    for ln in lines[:10]:
        if len(ln) >= 4 and len(ln) <= 60:
            # 排除明显是学号、日期、页眉等
            if re.search(r"\d{4,}", ln):
                continue
            if "目录" in ln or "Contents" in ln:
                continue
            cand.append(ln)
    if not cand:
        return lines[0] if lines else None
    # 优先包含报告/课程词的行
    for ln in cand:
        if any(x in ln for x in ["报告", "作业", "实验", "程序设计", "C语言"]):
            return ln
    # 否则返回最长
    return max(cand, key=len)

def extract_people(text: str) -> Dict[str, List[str]]:
    # 非严格：试着抓“组长/组员/姓名：xxx”
    people = {"组长": [], "组员": [], "姓名": []}
    patterns = {
        "组长": r"(组长|组\s*长)\s*[:：]\s*([^\n，,;；]{2,20})",
        "组员": r"(组员|组\s*员)\s*[:：]\s*([^\n]{2,60})",
        "姓名": r"(姓名)\s*[:：]\s*([^\n，,;；]{2,20})",
    }
    for k, pat in patterns.items():
        for m in re.finditer(pat, text):
            val = m.group(2).strip()
            val = re.sub(r"^[\|\s:：]+", "", val)
            val = re.sub(r"[\|\s]+$", "", val)
            val = re.sub(r"\s*\|\s*", " ", val)
            # 组员可能一行多个，用常见分隔符切一下
            if k == "组员":
                parts = re.split(r"[，,、\s]+", val)
                parts = [p for p in parts if 1 < len(p) < 20]
                people[k].extend(parts[:10])
            else:
                people[k].append(val)
    # 去重
    for k in people:
        seen = []
        for x in people[k]:
            if x not in seen:
                seen.append(x)
        people[k] = seen
    return people

def extract_topics(text: str) -> Tuple[List[str], Dict[str, List[str]]]:
    # 返回：topics + evidence（每个topic命中哪些关键词）
    found = []
    evidence: Dict[str, List[str]] = {}
    for topic, kws in TOPIC_KEYWORDS.items():
        hit_kws = [kw for kw in kws if kw in text]
        if hit_kws:
            found.append(topic)
            evidence[topic] = hit_kws[:8]
    # 可选：按命中数量排序，让“最相关”排前面
    found.sort(key=lambda t: len(evidence.get(t, [])), reverse=True)
    return found, evidence

@dataclass
class DocMeta:
    file: str
    ext: str
    doc_type: str
    title: Optional[str]
    course: Optional[str]
    week: Optional[int]
    topics: List[str]
    people: Dict[str, List[str]]
    evidence: Dict[str, object]

def extract_metadata(path: str) -> DocMeta:
    ext = os.path.splitext(path)[1].lower()
    lines: List[str] = []
    if ext == ".docx":
        lines = read_docx_text(path)
    elif ext == ".pdf":
        lines = read_pdf_text(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # 只用前面一部分做元数据识别更稳（避免正文噪声）
    head_text = "\n".join(lines[:120])
    all_text = "\n".join(lines)

    title = guess_title(lines)
    doc_type, type_hits = guess_doc_type(head_text)
    course = guess_course(head_text) or guess_course(all_text)
    week = guess_week(head_text) or guess_week(all_text)
    topics, topic_hits = extract_topics(all_text)
    people = extract_people(head_text)

    evidence = {
        "doc_type_hits": type_hits,
        "topic_hits": topic_hits,
        "sample_head": lines[:12],
    }

    return DocMeta(
        file=os.path.abspath(path),
        ext=ext,
        doc_type=doc_type,
        title=title,
        course=course,
        week=week,
        topics=topics[:8],  # 先截断，避免太长
        people=people,
        evidence=evidence,
    )

def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")):
                files.append(os.path.join(root, n))
    return sorted(files)

def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="metadata_results.json", help="output json filename")
    args = parser.parse_args()

    paths = iter_files(args.path)
    results = []
    for p in paths:
        try:
            meta = extract_metadata(p)
            results.append(asdict(meta))
            print(f"[OK] {os.path.basename(p)} -> type={meta.doc_type}, topics={meta.topics[:3]}")
        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {args.out} ({len(results)} items)")

if __name__ == "__main__":
    main()
