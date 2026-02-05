# question_extract_v2_2.py
import os
import re
import json
from dataclasses import dataclass, asdict
from typing import List, Optional, Tuple

# =====================
# 1) 文本读取：docx（含表格）/ pdf（保留）
# =====================

def read_docx_lines(path: str, max_paras: int = 1200, max_table_cells: int = 4000) -> List[str]:
    from docx import Document
    doc = Document(path)

    lines: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if len(lines) >= max_paras:
            break

    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = " ".join(
                    (pp.text or "").strip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()
                if cell_text:
                    row_parts.append(cell_text)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    cleaned = []
    for ln in lines:
        ln = ln.replace("\u3000", " ")
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned


def read_pdf_lines(path: str, max_pages: int = 8) -> List[str]:
    import fitz
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            ln = ln.replace("\u3000", " ")
            ln = re.sub(r"\s+", " ", ln).strip()
            if ln:
                lines.append(ln)
    return lines[:4000]


def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")


# =====================
# 2) 规则：疑问候选识别（v2.2）
# =====================

QUESTION_WORDS = [
    "为什么", "为何", "怎么", "如何", "怎样", "是否", "能否", "可否",
    "区别", "联系", "原因", "作用", "意义", "什么时候", "哪里", "多少",
    "是什么", "什么",  # v2.2：补“什么”以覆盖“...是什么/什么作用”
]

CONFUSION_WORDS = [
    "不懂", "没懂", "不理解", "没理解", "不太懂", "不太明白", "不清楚",
    "有点懵", "困惑", "疑惑", "看不懂", "搞不懂", "不确定", "不熟悉",
    "不会", "不知道", "不明白",
]

MARKER_WORDS = [
    "疑问", "问题", "困惑", "疑惑", "不懂", "不会", "没懂", "不理解", "不明白", "不清楚",
    "遇到的问题", "存在的问题",
]

# section
Q_SECTION_HINTS = [
    "自学心得或疑问", "心得或疑问", "心得与疑问", "反思与问题",
    "疑问", "问题", "困惑", "疑惑",
]
EXCLUDE_SECTION_HINTS = [
    "自学内容", "学习内容", "自学过程", "学习过程", "自学总结", "学习总结",
]
END_SECTION_HINTS = [
    "自学照片", "照片", "截图", "图片", "附件", "附录",
]

NEGATIVE_PATTERNS = [
    r"^目录$",
    r"^contents$",
    r"^第[一二三四五六七八九十0-9]{1,2}[章节]\b",
    r"^[0-9]+\s*$",
]

META_FIELD_HINTS = ["组长", "组员", "姓名", "学号", "班级", "日期", "学院", "专业"]

# v2.2：弱疑问提示词（用于切分/辅助，不直接当作强触发）
WEAK_QUESTION_HINTS = [
    "含义", "意思", "代表", "表示", "用途", "用法", "作用",
]

# v2.2：对“含义不明确/不清楚”这类弱表达，不直接当问题（除非同句有问号/疑问词/marker）
WEAK_CONFUSE_PAT = re.compile(r"(含义|意思|用法|用途|作用).{0,6}(不明确|不清楚|不太明白|不明白|不理解|不了解)")


@dataclass
class Candidate:
    source_line_no: int
    sent_no: int
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]


def is_negative_line(line: str) -> bool:
    l = line.strip().lower()
    for pat in NEGATIVE_PATTERNS:
        if re.search(pat, l, flags=re.IGNORECASE):
            return True
    return False


def looks_like_section_title(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if len(s) > 60:
        return False

    if re.match(r"^\s*(第?\s*[一二三四五六七八九十0-9]{1,2}\s*[、\.．:：\)])\s*", s):
        return True
    if re.match(r"^\s*(\(?\d{1,2}\)?\s*[、\.．:：\)])\s*", s):
        return True

    if any(h in s for h in (Q_SECTION_HINTS + EXCLUDE_SECTION_HINTS + END_SECTION_HINTS)):
        return True
    return False


def update_section_flag(line: str, in_q_section: bool) -> bool:
    s = line.strip()
    if any(h in s for h in END_SECTION_HINTS):
        return False
    if any(h in s for h in EXCLUDE_SECTION_HINTS):
        return False
    if any(h in s for h in Q_SECTION_HINTS):
        return True
    return in_q_section


# ---------- v2.2：分句，保留问号信号 ----------
def split_into_sentences_keep_q(line: str) -> List[str]:
    """
    与 v2 不同：切分时保留“？”（包括英文?统一为中文？）。
    """
    if not line:
        return []
    s = re.sub(r"\s+", " ", line).strip()
    if not s:
        return []

    # 统一问号
    s = s.replace("?", "？")

    # 先按句末标点切，但保留问号在子句里：用 finditer 自己切
    parts: List[str] = []
    start = 0
    for m in re.finditer(r"[。；;！!？]\s*", s):
        end = m.end()
        seg = s[start:end].strip()
        if seg:
            parts.append(seg)
        start = end
    last = s[start:].strip()
    if last:
        parts.append(last)

    # 对每个片段再按编号切（1. 2) (1)）
    final_parts: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        sub = re.split(r"(?:\s+|^)(?=(?:\(?\d{1,2}\)?[\.、\)])\s*)", p)
        sub = [x.strip() for x in sub if x.strip()]
        if len(sub) <= 1:
            final_parts.append(p)
        else:
            buf = ""
            for x in sub:
                if re.fullmatch(r"\(?\d{1,2}\)?[\.、\)]", x):
                    buf = x
                else:
                    if buf:
                        final_parts.append((buf + " " + x).strip())
                        buf = ""
                    else:
                        final_parts.append(x)
            if buf:
                final_parts.append(buf)

    final_parts = [x for x in final_parts if len(x) >= 2]
    return final_parts


# ---------- marker 拆分 ----------
ENUM_MARK_RE = re.compile(r"(?:^|\s)(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*")

def split_marker_tail(tail: str) -> List[str]:
    tail = tail.strip()
    if not tail:
        return []

    tail = tail.replace("?", "？")

    parts = split_into_sentences_keep_q(tail)
    if not parts:
        parts = [tail]

    refined: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue

        idxs = [m.start() for m in ENUM_MARK_RE.finditer(p)]
        if len(idxs) <= 1:
            refined.append(p)
            continue

        idxs.append(len(p))
        for a, b in zip(idxs, idxs[1:]):
            seg = p[a:b].strip()
            seg = re.sub(r"^\s*(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*", "", seg)
            seg = seg.strip()
            if seg:
                refined.append(seg)

    out = []
    for x in refined:
        x = re.sub(r"\s+", " ", x).strip()
        # 如果明显是问句但没问号，补一个（仅对含“什么/怎么/如何/是否”等强词）
        if "？" not in x and any(w in x for w in ["什么", "怎么", "如何", "为什么", "是否", "能否", "可否"]):
            x = x + "？"
        if len(x) >= 2:
            out.append(x)
    return out


def extract_marker_contents(text: str) -> List[str]:
    t = re.sub(r"\s*\|\s*", " | ", text)
    m = re.search(
        rf"({'|'.join(map(re.escape, MARKER_WORDS))})\s*(?:[:：]\s*|\|\s*[:：]?\s*\|\s*)(.+)$",
        t
    )
    if not m:
        return []

    tail = m.group(2).strip()
    tail = re.sub(r"^[\|\s:：]+", "", tail)
    tail = re.sub(r"\s*\|\s*", " ", tail).strip()
    if not tail:
        return []

    return split_marker_tail(tail)


def score_triggers(sent: str, in_q_section: bool) -> Tuple[List[str], int]:
    triggers: List[str] = []
    score = 0

    sent2 = sent.replace("?", "？")

    if "？" in sent2:
        triggers.append("HAS_QUESTION_MARK")
        score += 3

    for w in QUESTION_WORDS:
        if w in sent2:
            triggers.append(f"QWORD:{w}")
            score += 2
            break

    # 困惑词：在Q段更可信
    for w in CONFUSION_WORDS:
        if w in sent2:
            triggers.append(f"CONFUSE:{w}")
            score += 1 if not in_q_section else 2
            break

    # v2.2：弱疑问提示词，只加很小权重
    if any(h in sent2 for h in WEAK_QUESTION_HINTS):
        triggers.append("WEAK_HINT")
        score += 1 if in_q_section else 0

    if in_q_section and score > 0:
        score += 1
        triggers.append("IN_Q_SECTION")

    return triggers, score


def extract_candidates(lines: List[str]) -> List[Candidate]:
    cands: List[Candidate] = []
    in_q_section = False

    for i, raw_line in enumerate(lines):
        line = (raw_line or "").strip()
        if not line:
            continue
        if is_negative_line(line):
            continue

        # 标题更新 section
        if looks_like_section_title(line):
            in_q_section = update_section_flag(line, in_q_section)
            continue

        # 只在疑问段抽取（硬门控）
        if not in_q_section:
            continue

        prev_line = lines[i - 1].strip() if i - 1 >= 0 else None
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None

        # 元数据行过滤（但 marker / 问号保留）
        if any(f in line for f in META_FIELD_HINTS):
            has_qmark = ("?" in line or "？" in line)
            has_marker = bool(extract_marker_contents(line))
            if not has_qmark and not has_marker:
                continue

        # 策略A：显式 marker
        marker_parts = extract_marker_contents(line)
        if marker_parts:
            for j, p in enumerate(marker_parts):
                p = p.strip()
                if not p:
                    continue
                triggers, _ = score_triggers(p, in_q_section=True)
                triggers = (triggers or []) + ["MARKER:EXPLICIT", "IN_Q_SECTION"]
                cands.append(Candidate(
                    source_line_no=i,
                    sent_no=j,
                    text=p.replace("?", "？"),
                    triggers=sorted(list(set(triggers))),
                    in_q_section=True,
                    context_prev=prev_line,
                    context_next=next_line,
                ))
            continue

        # 策略B：分句
        sents = split_into_sentences_keep_q(line)
        for j, sent in enumerate(sents):
            sent = sent.strip()
            if not sent:
                continue

            sent = sent.replace("?", "？")

            # v2.2：如果句子包含“；”之类并且含弱困惑描述，优先把后面真正问句抽出
            # （例如：含义不明确 \t 制表位置是什么）
            # 这里不强行拆“含义不明确”为问题，避免误检
            if WEAK_CONFUSE_PAT.search(sent) and "？" not in sent and not any(w in sent for w in QUESTION_WORDS):
                # 弱困惑描述，不当作问题候选
                continue

            if len(sent) <= 4 and ("？" not in sent):
                continue

            triggers, score = score_triggers(sent, in_q_section=True)

            # 至少要 score>=2
            if score < 2:
                continue

            # 仅出现困惑词/弱提示但无问号无疑问词：过滤（避免“心得泛不懂”）
            only_weak = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                (any(t.startswith("CONFUSE:") for t in triggers) or "WEAK_HINT" in triggers)
            )
            if only_weak:
                continue

            cands.append(Candidate(
                source_line_no=i,
                sent_no=j,
                text=sent,
                triggers=sorted(list(set(triggers))),
                in_q_section=True,
                context_prev=prev_line,
                context_next=next_line,
            ))

    # 去重
    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        key = c.text
        if key not in seen:
            uniq.append(c)
            seen.add(key)

    return uniq


# =====================
# 3) 批处理：文件/文件夹
# =====================

def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")) and not n.startswith("~$"):
                files.append(os.path.join(root, n))
    return sorted(files)


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="question_candidates_v2_2.json", help="output json filename")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file for quick check")
    args = parser.parse_args()

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)
            cands = extract_candidates(lines)

            item = {
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(cands),
                "candidates": [asdict(x) for x in cands],
            }
            results.append(item)

            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(cands)}")

            if args.max_preview > 0:
                for x in cands[:args.max_preview]:
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) {x.text}  triggers={x.triggers}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {args.out} ({len(results)} files)")


if __name__ == "__main__":
    main()
