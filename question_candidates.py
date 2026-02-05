import os
import re
import json
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple

# =====================
# 1) 文本读取：docx / pdf（含表格）
# =====================

def read_docx_lines(path: str, max_paras: int = 800, max_table_cells: int = 2000) -> List[str]:
    from docx import Document
    doc = Document(path)

    lines: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if len(lines) >= max_paras:
            break

    # tables
    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = " ".join(
                    (pp.text or "").strip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()
                if cell_text:
                    row_parts.append(cell_text)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    # normalize
    cleaned = []
    for ln in lines:
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned


def read_pdf_lines(path: str, max_pages: int = 8) -> List[str]:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            ln = re.sub(r"\s+", " ", ln).strip()
            if ln:
                lines.append(ln)
    return lines[:4000]


def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")


# =====================
# 2) 规则：疑问候选识别（v2）
# =====================

QUESTION_WORDS = [
    "为什么", "为何", "怎么", "如何", "怎样", "是否", "能否", "可否",
    "区别", "联系", "原因", "作用", "意义", "什么时候", "哪里", "多少",
    "是什么",
]

CONFUSION_WORDS = [
    "不懂", "没懂", "不理解", "没理解", "不太懂", "不太明白", "不清楚",
    "有点懵", "困惑", "疑惑", "看不懂", "搞不懂", "不确定", "不熟悉",
    "不会", "不知道", "不明白",
]

# 显式“标签”词：不仅仅是“疑问”
MARKER_WORDS = [
    "疑问", "问题", "困惑", "疑惑", "不懂", "不会", "没懂", "不理解", "不明白", "不清楚",
    "遇到的问题", "存在的问题",
]

# 章节标题：进入这些章节后，可以适当放宽（提升召回）
SECTION_HINTS = [
    "疑问", "问题", "困惑", "疑惑",
    "自学心得或疑问", "心得或疑问", "心得与疑问", "反思与问题",
]

NEGATIVE_PATTERNS = [
    r"^目录$",
    r"^contents$",
    r"^第[一二三四五六七八九十0-9]{1,2}[章节]\b",
    r"^[0-9]+\s*$",
]

META_FIELD_HINTS = ["组长", "组员", "姓名", "学号", "班级", "日期", "学院", "专业"]


@dataclass
class Candidate:
    source_line_no: int          # 原始行号
    sent_no: int                 # 该行拆分后的子句序号
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]


def is_negative_line(line: str) -> bool:
    l = line.strip().lower()
    for pat in NEGATIVE_PATTERNS:
        if re.search(pat, l, flags=re.IGNORECASE):
            return True
    return False


def looks_like_section_title(line: str) -> bool:
    # 简单判断：短、含章节提示词、像标题
    if len(line) > 50:
        return False
    # 常见编号：一、二、1. 2) 等
    if re.match(r"^\s*([一二三四五六七八九十0-9]+[、\.．\)])\s*", line):
        pass
    # 包含提示词
    return any(h in line for h in SECTION_HINTS)


def split_into_sentences(line: str) -> List[str]:
    """
    v2 核心：把“长行”切分成更细颗粒的子句。
    同时兼容：中文标点、分号、编号（1. 2)）、以及 docx 合并导致的长段。
    """
    s = re.sub(r"\s+", " ", line).strip()
    if not s:
        return []

    # 先用常见句末标点切
    parts = re.split(r"[。！？?!；;]\s*", s)

    # 对每个片段再按编号切（如：1. xxx 2. yyy）
    final_parts: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 按 “ 1. / 1、 / 1) / (1) ” 这类编号切分
        sub = re.split(r"(?:\s+|^)(?=(?:\(?\d{1,2}\)?[\.、\)])\s*)", p)
        # 上面 split 可能产生空串，清理一下
        sub = [x.strip() for x in sub if x.strip()]
        if len(sub) <= 1:
            final_parts.append(p)
        else:
            # 还原编号粘连问题：如果某段只有编号，把它和下一段拼起来
            buf = ""
            for x in sub:
                if re.fullmatch(r"\(?\d{1,2}\)?[\.、\)]", x):
                    buf = x
                else:
                    if buf:
                        final_parts.append((buf + " " + x).strip())
                        buf = ""
                    else:
                        final_parts.append(x)
            if buf:
                final_parts.append(buf)

    # 去掉过短噪声
    final_parts = [x for x in final_parts if len(x) >= 2]
    return final_parts


def extract_marker_contents(text: str) -> List[str]:
    """
    匹配 “标签：内容” 的结构，并返回标签后的内容（可能含多条问题）。
    例如：
      疑问：1.xxx 2.yyy
      问题: 为什么...
    """
    # 统一 | 分隔符，避免表格把 : 拆开
    t = re.sub(r"\s*\|\s*", " | ", text)

    # 允许： 疑问 ： | 内容  或 疑问:内容
    # 捕获 marker 和后续内容
    m = re.search(rf"({'|'.join(map(re.escape, MARKER_WORDS))})\s*(?:[:：]\s*|\|\s*[:：]?\s*\|\s*)(.+)$", t)
    if not m:
        return []

    tail = m.group(2).strip()
    # 清掉可能残留的表格符号
    tail = re.sub(r"^[\|\s:：]+", "", tail)
    tail = re.sub(r"\s*\|\s*", " ", tail).strip()

    if not tail:
        return []

    # 多问题拆分：按 “； ; / 换段 / 编号” 等
    parts = split_into_sentences(tail)

    # 如果 split 不开（比如只有一条），也返回
    return parts if parts else [tail]


def score_triggers(sent: str, in_q_section: bool) -> Tuple[List[str], int]:
    """
    为子句打触发规则，并给一个粗评分（用于过滤太弱的候选）。
    """
    triggers: List[str] = []
    score = 0

    # 问号
    if "?" in sent or "？" in sent:
        triggers.append("HAS_QUESTION_MARK")
        score += 3

    # 疑问词
    for w in QUESTION_WORDS:
        if w in sent:
            triggers.append(f"QWORD:{w}")
            score += 2
            break

    # 困惑词
    for w in CONFUSION_WORDS:
        if w in sent:
            triggers.append(f"CONFUSE:{w}")
            score += 2
            break

    # 在“疑问章节”中，允许更弱信号（例如没问号但提到“不理解/困惑”）
    if in_q_section and score > 0:
        score += 1
        triggers.append("IN_Q_SECTION")

    return triggers, score


def extract_candidates(lines: List[str]) -> List[Candidate]:
    cands: List[Candidate] = []
    in_q_section = False

    for i, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            continue
        if is_negative_line(line):
            continue

        # 检测章节切换（遇到“疑问/问题”标题就进入疑问章节）
        if looks_like_section_title(line):
            if any(h in line for h in SECTION_HINTS):
                in_q_section = True

        # 元数据行通常不当问题（除非真的带问号或显式marker）
        if any(f in line for f in META_FIELD_HINTS):
            has_qmark = ("?" in line or "？" in line)
            has_marker = bool(extract_marker_contents(line))
            if not has_qmark and not has_marker:
                continue

        prev_line = lines[i - 1].strip() if i - 1 >= 0 else None
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None

        # --- 策略A：显式 marker（高精度）---
        marker_parts = extract_marker_contents(line)
        if marker_parts:
            for j, p in enumerate(marker_parts):
                p = p.strip()
                if not p:
                    continue
                # marker 抽到的内容默认算候选，仍然打触发标签
                triggers, score = score_triggers(p, in_q_section=True)
                triggers = (triggers or []) + ["MARKER:EXPLICIT"]
                # 只要 marker 命中就不过分过滤
                cands.append(Candidate(
                    source_line_no=i,
                    sent_no=j,
                    text=p,
                    triggers=sorted(list(set(triggers))),
                    in_q_section=True,
                    context_prev=prev_line,
                    context_next=next_line,
                ))
            continue  # marker 行不再走下面策略，避免重复/粘连

        # --- 策略B：分句后用通用规则 ---
        sents = split_into_sentences(line)
        for j, sent in enumerate(sents):
            sent = sent.strip()
            if not sent:
                continue
            # 过滤特别短的片段（除非有问号）
            if len(sent) <= 4 and ("?" not in sent and "？" not in sent):
                continue

            triggers, score = score_triggers(sent, in_q_section=in_q_section)

            # v2 过滤策略：
            # - 正常情况下，至少要 score>=3（比如问号，或疑问词+困惑词）
            # - 如果处于疑问章节，score>=2 也可接受
            if in_q_section:
                if score < 2:
                    continue
            else:
                if score < 3:
                    continue

            cands.append(Candidate(
                source_line_no=i,
                sent_no=j,
                text=sent,
                triggers=sorted(list(set(triggers))),
                in_q_section=in_q_section,
                context_prev=prev_line,
                context_next=next_line,
            ))

    # 去重：相同文本重复出现（页眉/页脚/表格重复）
    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        key = c.text
        if key not in seen:
            uniq.append(c)
            seen.add(key)

    return uniq


# =====================
# 3) 批处理：文件/文件夹
# =====================

def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")):
                files.append(os.path.join(root, n))
    return sorted(files)


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="question_candidates_v2.json", help="output json filename")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file for quick check")
    args = parser.parse_args()

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)
            cands = extract_candidates(lines)

            item = {
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(cands),
                "candidates": [asdict(x) for x in cands],
            }
            results.append(item)

            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(cands)}")

            if args.max_preview > 0:
                for x in cands[:args.max_preview]:
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) {x.text}  triggers={x.triggers}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {args.out} ({len(results)} files)")

if __name__ == "__main__":
    main()
