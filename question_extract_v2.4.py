#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
question_extract_v2.4.py  (RULE-ONLY)

主流程：
1) 用 v2.4 规则抽取疑问候选（疑问段硬门控）
2) 不做 LLM 补漏（num_candidates==0 就是 0）
3) 输出 JSON 结构对齐 question_candidates_v2_3.json
4) 输出文件自动避免重名覆盖：若已存在则追加时间戳

v2.4 规则特性（与 v2_3_fallback_llm_fixed_full 抽取规则一致）：
- DOCX 读取保留段落内换行（splitlines）
- 支持 “问题/疑问/困惑/疑惑” 单独一行标题 -> 连续吸收后续枚举条目（允许名词化）
- 当检测到“显式 marker 行”（如 “问题：1.xxx”）时：
  A) 先抽取该行尾部 marker_parts
  B) 再向后吸收紧随其后的连续编号行（2.xxx / 3.xxx ...）
- 避免“心得：...区别...”误抽取为问题（心得行若无问号/困惑词/marker则跳过）

运行示例：
python question_extract_v2.4.py norm_dataset --out out.json

调试某文件：
python question_extract_v2.4.py norm_dataset --debug_file 符丽 --max_preview 50
"""

import os
import re
import json
import argparse
from dataclasses import dataclass, asdict
from typing import List, Optional, Tuple, Dict, Any


# =====================
# 0) 工具：输出避免重名
# =====================
def make_nonconflicting_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if not ext:
        ext = ".json"
    return f"{base}_{ts}{ext}"


# =====================
# 1) 文本读取：docx/pdf
# =====================
def _clean_lines(lines: List[str]) -> List[str]:
    """
    不要用 \\s+ 全压缩（会破坏结构），这里只压缩连续空格/制表，保留“行”的概念。
    """
    cleaned = []
    for ln in lines:
        ln = (ln or "").replace("\u3000", " ").strip()
        ln = re.sub(r"[ \t]+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned


def read_docx_lines(path: str, max_paras: int = 1200, max_table_cells: int = 4000) -> List[str]:
    """
    关键：保留段落内部换行（splitlines），避免“问题：1... / 2...”被压扁到一行。
    """
    from docx import Document
    doc = Document(path)

    lines: List[str] = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "")
        for seg in t.splitlines():  # 保留段落内换行
            seg = (seg or "").strip()
            if seg:
                lines.append(seg)
        if len(lines) >= max_paras:
            break

    # 表格
    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = "\n".join(
                    (pp.text or "").rstrip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()
                if cell_text:
                    for seg in cell_text.splitlines():
                        seg = seg.strip()
                        if seg:
                            row_parts.append(seg)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    return _clean_lines(lines)


def read_pdf_lines(path: str, max_pages: int = 8) -> List[str]:
    import fitz
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            ln = (ln or "").replace("\u3000", " ").strip()
            ln = re.sub(r"[ \t]+", " ", ln).strip()
            if ln:
                lines.append(ln)
    return lines[:4000]


def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")


def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")) and not n.startswith("~$"):
                files.append(os.path.join(root, n))
    return sorted(files)


# =====================
# 2) v2.4 规则抽取（与 v2_3_fixed_full 抽取规则保持一致）
# =====================
QUESTION_WORDS = [
    "为什么", "为何", "怎么", "如何", "怎样", "是否", "能否", "可否",
    "联系", "原因", "作用", "意义", "什么时候", "哪里", "多少",
    "是什么", "什么", "吗"
]

CONFUSION_WORDS = [
    "不懂", "没懂", "不理解", "没理解", "不太懂", "不太明白", "不清楚",
    "有点懵", "困惑", "疑惑", "看不懂", "搞不懂", "不确定", "不熟悉",
    "不会", "不知道", "不明白", "搞混", "不了解",
    "把握不好", "掌握不好", "卡壳"
]

MARKER_WORDS = [
    "疑问", "问题", "困惑", "疑惑", "不懂", "不会", "没懂", "不理解", "不明白", "不清楚",
    "遇到的问题", "存在的问题", "有困难"
]

Q_SECTION_HINTS = [
    "自学心得或疑问", "心得或疑问", "心得与疑问", "反思与问题",
    "疑问", "问题", "困惑", "疑惑",
]

EXCLUDE_SECTION_HINTS = [
    "自学内容", "学习内容", "自学过程", "学习过程", "自学总结", "学习总结",
]

END_SECTION_HINTS = [
    "自学照片", "照片", "截图", "图片", "附件", "附录",
]

NEGATIVE_PATTERNS = [
    r"^目录$",
    r"^contents$",
    r"^第[一二三四五六七八九十0-9]{1,2}[章节]\b",
    r"^[0-9]+\s*$",
]

META_FIELD_HINTS = ["组长", "组员", "姓名", "学号", "班级", "日期", "学院", "专业"]

WEAK_QUESTION_HINTS = [
    "含义", "意思", "代表", "表示", "用途", "用法", "作用",
]

WEAK_CONFUSE_PAT = re.compile(r"(含义|意思|用法|用途|作用).{0,6}(不明确|不清楚|不太明白|不明白|不理解|不了解)")

COURSE_SIGNAL_WORDS = [
    "for", "while", "do", "if", "else", "switch", "case", "break", "continue",
    "scanf", "printf", "gets", "puts", "fopen", "fclose", "fprintf", "fscanf",
    "%d", "%f", "%c", "%s", "\\t", "\\n", "\\0",
    "数组", "字符串", "字符数组", "指针", "函数", "递归", "循环", "分支", "进制",
    "变量", "常量", "数据类型", "结构体", "链表", "文件", "输入", "输出",
    "ASCII", "float", "double", "int", "char"
]

# 枚举条目识别：1. / 1、 / (1) / 一、 / 一. ...
ENUM_LINE_RE = re.compile(
    r"^\s*(?:\(?\d{1,3}\)?[\.、\)]|[一二三四五六七八九十]{1,3}[、\.．\)])\s*"
)


def is_enum_line(line: str) -> bool:
    return bool(ENUM_LINE_RE.match((line or "").strip()))


def strip_enum_prefix(line: str) -> str:
    return re.sub(ENUM_LINE_RE, "", (line or "")).strip()


def is_marker_header_line(line: str) -> bool:
    """
    识别：单独一行的 ‘疑问/问题/困惑/疑惑’（可带冒号）
    """
    s = (line or "").strip()
    if not s:
        return False
    s2 = re.sub(r"\s*\|\s*", " ", s).strip()
    s2n = s2.replace(" ", "")
    s2n = re.sub(r"[：:]+$", "", s2n)
    return s2n in ("疑问", "问题", "困惑", "疑惑")


# 避免“心得：...”误抽取
XINDE_PREFIX_RE = re.compile(r"^\s*心得\s*[:：]")


@dataclass
class Candidate:
    source_line_no: int
    sent_no: int
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]
    llm: Optional[Dict[str, Any]] = None


def is_negative_line(line: str) -> bool:
    l = line.strip().lower()
    for pat in NEGATIVE_PATTERNS:
        if re.search(pat, l, flags=re.IGNORECASE):
            return True
    return False


def looks_like_section_title(line: str) -> bool:
    """
    v2.4：尽量不要把“2.xxx”当标题（否则会影响续行枚举）
    我们只把这些算标题：
      - 明确的段落提示词（Q_SECTION_HINTS / EXCLUDE_SECTION_HINTS / END_SECTION_HINTS）
      - 第X章/节
    不再把“纯数字编号开头”当标题（避免误伤 2.xxx）
    """
    s = line.strip()
    if not s:
        return False
    if "？" in s or any(w in s for w in ["如何", "怎么", "为什么", "是什么", "是否", "能否", "可否", "吗"]):
        return False

    title_like = s.replace("：", "").replace(":", "").strip()
    if title_like in Q_SECTION_HINTS or title_like in EXCLUDE_SECTION_HINTS or title_like in END_SECTION_HINTS:
        return True

    # 章节标题
    if re.match(r"^\s*(第?\s*[一二三四五六七八九十0-9]{1,2}\s*[、\.．:：\)])\s*", s):
        return True

    return False


def update_section_flag(line: str, in_q_section: bool) -> bool:
    s = line.strip()
    if any(h in s for h in END_SECTION_HINTS):
        return False
    if any(h in s for h in EXCLUDE_SECTION_HINTS):
        return False
    if any(h in s for h in Q_SECTION_HINTS):
        return True
    return in_q_section


def split_into_sentences_keep_q(line: str) -> List[str]:
    if not line:
        return []
    s = re.sub(r"[ \t]+", " ", line).strip()
    if not s:
        return []
    s = s.replace("?", "？")

    parts: List[str] = []
    start = 0
    for m in re.finditer(r"[。；;！!？]\s*", s):
        end = m.end()
        seg = s[start:end].strip()
        if seg:
            parts.append(seg)
        start = end
    last = s[start:].strip()
    if last:
        parts.append(last)

    # 轻度编号切分（行内的 1. xxx 2. yyy）
    final_parts: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        sub = re.split(r"(?:\s+|^)(?=(?:\(?\d{1,2}\)?[\.、\)])\s*)", p)
        sub = [x.strip() for x in sub if x.strip()]
        if len(sub) <= 1:
            final_parts.append(p)
        else:
            buf = ""
            for x in sub:
                if re.fullmatch(r"\(?\d{1,2}\)?[\.、\)]", x):
                    buf = x
                else:
                    if buf:
                        final_parts.append((buf + " " + x).strip())
                        buf = ""
                    else:
                        final_parts.append(x)
            if buf:
                final_parts.append(buf)

    return [x for x in final_parts if len(x) >= 2]


ENUM_MARK_RE = re.compile(r"(?:^|\s)(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*")


def split_marker_tail(tail: str) -> List[str]:
    tail = tail.strip()
    if not tail:
        return []
    tail = tail.replace("?", "？")

    rough_parts: List[str] = []
    for seg in tail.splitlines():
        seg = seg.strip()
        if not seg:
            continue
        rough_parts.extend(split_into_sentences_keep_q(seg) or [seg])

    parts = [x for x in rough_parts if x.strip()]

    refined: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue

        idxs = [m.start() for m in ENUM_MARK_RE.finditer(p)]
        if len(idxs) <= 1:
            refined.append(p)
            continue

        idxs.append(len(p))
        for a, b in zip(idxs, idxs[1:]):
            seg = p[a:b].strip()
            seg = re.sub(
                r"^\s*(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*",
                "",
                seg
            )
            seg = seg.strip()
            if seg:
                refined.append(seg)

    out = []
    for x in refined:
        x = re.sub(r"[ \t]+", " ", x).strip()
        if "？" not in x and any(w in x for w in ["什么", "怎么", "如何", "为什么", "是否", "能否", "可否"]):
            x = x + "？"
        if len(x) >= 2:
            out.append(x)
    return out


def extract_marker_contents(text: str) -> List[str]:
    t = re.sub(r"\s*\|\s*", " | ", text)
    m = re.search(
        rf"({'|'.join(map(re.escape, MARKER_WORDS))})\s*(?:[:：]\s*|\|\s*[:：]?\s*\|\s*)(.+)$",
        t
    )
    if not m:
        return []

    tail = m.group(2).strip()
    tail = re.sub(r"^[\|\s:：]+", "", tail)
    tail = re.sub(r"\s*\|\s*", " ", tail).strip()
    if not tail:
        return []

    return split_marker_tail(tail)


def score_triggers(sent: str, in_q_section: bool) -> Tuple[List[str], int]:
    triggers: List[str] = []
    score = 0
    sent2 = sent.replace("?", "？")

    if "？" in sent2:
        triggers.append("HAS_QUESTION_MARK")
        score += 3

    for w in QUESTION_WORDS:
        if w in sent2:
            triggers.append(f"QWORD:{w}")
            score += 2
            break

    for w in CONFUSION_WORDS:
        if w in sent2:
            triggers.append(f"CONFUSE:{w}")
            score += 2 if not in_q_section else 3
            break

    if any(h in sent2 for h in WEAK_QUESTION_HINTS):
        triggers.append("WEAK_HINT")
        score += 1 if in_q_section else 0

    if in_q_section and score > 0:
        score += 1
        triggers.append("IN_Q_SECTION")

    return triggers, score


def extract_candidates_v2_4(lines: List[str]) -> List[Candidate]:
    cands: List[Candidate] = []
    in_q_section = False

    # 续行枚举模式：遇到“问题/疑问/困惑/疑惑”标题 或 “问题：xxx”显式marker后可开启
    marker_cont_active = False

    for i, raw_line in enumerate(lines):
        line = (raw_line or "").strip()
        if not line:
            marker_cont_active = False
            continue
        if is_negative_line(line):
            marker_cont_active = False
            continue

        prev_line = lines[i - 1].strip() if i - 1 >= 0 else None
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None

        # 标题更新 section
        if looks_like_section_title(line):
            in_q_section = update_section_flag(line, in_q_section)
            marker_cont_active = False
            continue

        # 只在疑问段抽取（硬门控）
        if not in_q_section:
            marker_cont_active = False
            continue

        # 避免“心得：...区别...”误抽取
        if XINDE_PREFIX_RE.search(line):
            tmp_triggers, _ = score_triggers(line, in_q_section=True)
            has_explicit_marker = bool(extract_marker_contents(line))
            has_confuse = any(t.startswith("CONFUSE:") for t in tmp_triggers)
            has_q = ("HAS_QUESTION_MARK" in tmp_triggers) or any(t.startswith("QWORD:") for t in tmp_triggers)
            if (not has_explicit_marker) and (not has_confuse) and (not has_q):
                continue

        # 元数据行过滤（但 marker / 问号保留）
        if any(f in line for f in META_FIELD_HINTS):
            has_qmark = ("?" in line or "？" in line)
            has_marker = bool(extract_marker_contents(line))
            if not has_qmark and not has_marker:
                continue

        # “疑问/问题/困惑/疑惑” 单独一行标题 -> 开启续行吸收
        if is_marker_header_line(line):
            marker_cont_active = True
            continue

        # -------- 显式 marker（同一行：问题：xxx）--------
        marker_parts = extract_marker_contents(line)
        if marker_parts:
            # 1) 本行 marker 解析出来的内容加入候选
            for j, ptxt in enumerate(marker_parts):
                ptxt = (ptxt or "").strip()
                if not ptxt:
                    continue
                triggers, _ = score_triggers(ptxt, in_q_section=True)
                triggers = (triggers or []) + ["MARKER:EXPLICIT", "IN_Q_SECTION"]
                cands.append(Candidate(
                    source_line_no=i,
                    sent_no=j,
                    text=ptxt.replace("?", "？"),
                    triggers=sorted(list(set(triggers))),
                    in_q_section=True,
                    context_prev=prev_line,
                    context_next=next_line,
                    llm=None
                ))

            # 2) 关键修复：向后吸收连续编号行（2.xxx / 3.xxx ...）
            k = i + 1
            sent_no = len(marker_parts)
            while k < len(lines):
                nxt = (lines[k] or "").strip()
                if not nxt:
                    break
                if any(h in nxt for h in END_SECTION_HINTS):
                    break
                if not is_enum_line(nxt):
                    break

                body = strip_enum_prefix(nxt).replace("?", "？").strip()
                if body:
                    triggers, _ = score_triggers(body, in_q_section=True)
                    triggers = (triggers or []) + ["MARKER:FOLLOWING_ENUM", "IN_Q_SECTION"]
                    cands.append(Candidate(
                        source_line_no=k,
                        sent_no=sent_no,
                        text=body,
                        triggers=sorted(list(set(triggers))),
                        in_q_section=True,
                        context_prev=lines[k - 1].strip() if k - 1 >= 0 else None,
                        context_next=lines[k + 1].strip() if k + 1 < len(lines) else None,
                        llm=None
                    ))
                    sent_no += 1

                k += 1

            marker_cont_active = True
            continue

        # -------- marker 续行枚举条目（允许名词化）--------
        if marker_cont_active and is_enum_line(line):
            body = strip_enum_prefix(line).replace("?", "？").strip()
            if body and len(body) >= 2:
                triggers, _ = score_triggers(body, in_q_section=True)
                triggers = (triggers or []) + ["MARKER:CONT", "IN_Q_SECTION"]
                if len(body) >= 3:
                    cands.append(Candidate(
                        source_line_no=i,
                        sent_no=0,
                        text=body,
                        triggers=sorted(list(set(triggers))),
                        in_q_section=True,
                        context_prev=prev_line,
                        context_next=next_line,
                        llm=None
                    ))
            continue

        # marker 续行遇到非枚举行：若无疑问信号则关闭
        if marker_cont_active:
            tmp_triggers, _ = score_triggers(line.replace("?", "？"), in_q_section=True)
            has_signal = (
                ("HAS_QUESTION_MARK" in tmp_triggers)
                or any(t.startswith("QWORD:") for t in tmp_triggers)
                or any(t.startswith("CONFUSE:") for t in tmp_triggers)
            )
            if not has_signal:
                marker_cont_active = False

        # -------- 原有：分句抽取 --------
        sents = split_into_sentences_keep_q(line)
        for j, sent in enumerate(sents):
            sent = sent.strip()
            if not sent:
                continue
            sent = sent.replace("?", "？")

            if WEAK_CONFUSE_PAT.search(sent) and "？" not in sent and not any(w in sent for w in QUESTION_WORDS):
                continue

            if len(sent) <= 4 and ("？" not in sent):
                continue

            triggers, score = score_triggers(sent, in_q_section=True)
            if score < 2:
                continue

            only_weak = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                (any(t.startswith("CONFUSE:") for t in triggers) or "WEAK_HINT" in triggers)
            )

            confuse_only = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                any(t.startswith("CONFUSE:") for t in triggers)
            )

            sent_l = sent.lower()
            has_course_signal = any(k.lower() in sent_l for k in COURSE_SIGNAL_WORDS)

            if only_weak:
                if confuse_only and has_course_signal and len(sent) >= 8:
                    triggers.append("ALLOW:CONFUSE_WITH_SIGNAL")
                else:
                    continue

            cands.append(Candidate(
                source_line_no=i,
                sent_no=j,
                text=sent,
                triggers=sorted(list(set(triggers))),
                in_q_section=True,
                context_prev=prev_line,
                context_next=next_line,
                llm=None
            ))

    # 去重
    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        if c.text not in seen:
            uniq.append(c)
            seen.add(c.text)
    return uniq


# =====================
# 3) 主流程：纯规则
# =====================
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="question_candidates_v2_4_rule_only.json", help="output json filename")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file for quick check")

    # debug
    parser.add_argument("--debug_file", default="", help="只调试打印包含该关键词的文件名（如 符丽）")
    parser.add_argument("--debug_kw", default="问题", help="调试时定位关键词（默认：问题）")
    args = parser.parse_args()

    out_path = make_nonconflicting_path(args.out)

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)

            # ===== DEBUG PRINT =====
            if args.debug_file and (args.debug_file in os.path.basename(p)):
                print("\n" + "=" * 80)
                print("[DEBUG] file:", p)
                print("[DEBUG] total lines:", len(lines))
                print("\n[DEBUG] first 80 lines (repr):")
                for idx, ln in enumerate(lines[:80]):
                    print(f"{idx:03d}: {repr(ln)}")

                hits = [i for i, ln in enumerate(lines) if args.debug_kw in ln]
                print(f"\n[DEBUG] hit indices for kw='{args.debug_kw}':", hits)

                for hi in hits[:10]:
                    print("\n[DEBUG] around index", hi)
                    for j in range(max(0, hi - 5), min(len(lines), hi + 10)):
                        print(f"{j:03d}: {repr(lines[j])}")
                print("=" * 80 + "\n")
            # ===== DEBUG PRINT END =====

            rule_cands = extract_candidates_v2_4(lines)

            results.append({
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(rule_cands),
                "candidates": [asdict(x) for x in rule_cands],
                "fallback_used": False,  # 纯规则版固定为 False
            })

            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(rule_cands)} mode=RULE_ONLY")

            if args.max_preview > 0:
                for x in rule_cands[:args.max_preview]:
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) {x.text} triggers={x.triggers}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {out_path} ({len(results)} files)")


if __name__ == "__main__":
    main()
