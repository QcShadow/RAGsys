#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
question_extract_v2_3_fallback_llm_markercontv3.py

主流程：
1) 先用 v2.3 规则抽取疑问候选
2) 若某文件规则抽取 num_candidates == 0，则调用 strict prompt 的 LLM 做兜底抽取
   - LLM 允许输出 []（表示确实没有可用于RAG问答的疑问/困惑点）
3) 输出 JSON 结构对齐 question_candidates_v2_3.json
4) 输出文件自动避免重名覆盖：若已存在则追加时间戳

本版修复：
A) DOCX 读取保留段落内换行（splitlines），避免“问题：1... / 2...”被压扁
B) 支持 “问题/疑问/困惑/疑惑” 单独一行标题 -> 连续吸收后续枚举条目（允许名词化）
C) 避免“心得：...区别...”误抽取为问题（心得行若无问号/困惑词/marker则跳过）
"""

import os
import re
import json
import argparse
from dataclasses import dataclass, asdict
from typing import List, Optional, Tuple, Dict, Any

# ---------------------
# LLM 配置：建议用环境变量覆盖
# ---------------------
OPENAI_API_KEY="sk-66df7e992c384ffd8d352fb948465d2e"
OPENAI_BASE_URL="https://api.deepseek.com"
OPENAI_MODEL="deepseek-chat"
OPENAI_TIMEOUT=120

# =====================
# 0) 工具：输出避免重名
# =====================

def make_nonconflicting_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if not ext:
        ext = ".json"
    return f"{base}_{ts}{ext}"

# =====================
# 1) 文本读取：docx/pdf
# =====================

def _clean_lines(lines: List[str]) -> List[str]:
    """
    关键：不要用 \\s+ 全压缩（会把结构性换行等“变形”后的遗留符号也干掉）
    我们在上游已经 splitlines()，这里只压缩连续空格/制表，保留“行”的概念。
    """
    cleaned = []
    for ln in lines:
        ln = (ln or "").replace("\u3000", " ").strip()
        ln = re.sub(r"[ \t]+", " ", ln).strip()
        if ln and len(ln) >= 2:
            cleaned.append(ln)
    return cleaned

def read_docx_lines(path: str, max_paras: int = 1200, max_table_cells: int = 4000) -> List[str]:
    """
    关键修复：保留段落内部换行（splitlines），避免“问题：1... / 2...”被压扁到一行。
    """
    from docx import Document
    doc = Document(path)

    lines: List[str] = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "")
        # 保留段落内部的手动换行
        for seg in t.splitlines():
            seg = (seg or "").strip()
            if seg:
                lines.append(seg)
        if len(lines) >= max_paras:
            break

    # 表格（每个 cell 也 splitlines）
    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = "\n".join(
                    (pp.text or "").rstrip() for pp in cell.paragraphs if (pp.text or "").strip()
                ).strip()

                if cell_text:
                    for seg in cell_text.splitlines():
                        seg = seg.strip()
                        if seg:
                            row_parts.append(seg)

                cell_count += 1
                if cell_count >= max_table_cells:
                    break

            row_text = " | ".join(row_parts).strip()
            if row_text:
                lines.append(row_text)

            if cell_count >= max_table_cells:
                break
        if cell_count >= max_table_cells:
            break

    return _clean_lines(lines)

def read_pdf_lines(path: str, max_pages: int = 8) -> List[str]:
    import fitz
    doc = fitz.open(path)
    lines: List[str] = []
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        for ln in text.splitlines():
            ln = (ln or "").replace("\u3000", " ").strip()
            ln = re.sub(r"[ \t]+", " ", ln).strip()
            if ln:
                lines.append(ln)
    return lines[:4000]

def read_lines(path: str) -> List[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_lines(path)
    if ext == ".pdf":
        return read_pdf_lines(path)
    raise ValueError(f"Unsupported file: {path}")

def iter_files(input_path: str) -> List[str]:
    if os.path.isfile(input_path):
        return [input_path]
    files = []
    for root, _, names in os.walk(input_path):
        for n in names:
            if n.lower().endswith((".docx", ".pdf")) and not n.startswith("~$"):
                files.append(os.path.join(root, n))
    return sorted(files)

# =====================
# 2) v2.3 规则抽取（内置）
# =====================

QUESTION_WORDS = [
    "为什么", "为何", "怎么", "如何", "怎样", "是否", "能否", "可否",
    "区别", "联系", "原因", "作用", "意义", "什么时候", "哪里", "多少",
    "是什么", "什么", "吗"
]

CONFUSION_WORDS = [
    "不懂", "没懂", "不理解", "没理解", "不太懂", "不太明白", "不清楚",
    "有点懵", "困惑", "疑惑", "看不懂", "搞不懂", "不确定", "不熟悉",
    "不会", "不知道", "不明白", "搞混", "不了解",
    "把握不好", "掌握不好", "卡壳"
]

MARKER_WORDS = [
    "疑问", "问题", "困惑", "疑惑", "不懂", "不会", "没懂", "不理解", "不明白", "不清楚",
    "遇到的问题", "存在的问题", "有困难"
]

Q_SECTION_HINTS = [
    "自学心得或疑问", "心得或疑问", "心得与疑问", "反思与问题",
    "疑问", "问题", "困惑", "疑惑",
]
EXCLUDE_SECTION_HINTS = [
    "自学内容", "学习内容", "自学过程", "学习过程", "自学总结", "学习总结",
]
END_SECTION_HINTS = [
    "自学照片", "照片", "截图", "图片", "附件", "附录",
]

NEGATIVE_PATTERNS = [
    r"^目录$",
    r"^contents$",
    r"^第[一二三四五六七八九十0-9]{1,2}[章节]\b",
    r"^[0-9]+\s*$",
]

META_FIELD_HINTS = ["组长", "组员", "姓名", "学号", "班级", "日期", "学院", "专业"]

WEAK_QUESTION_HINTS = [
    "含义", "意思", "代表", "表示", "用途", "用法", "作用",
]

WEAK_CONFUSE_PAT = re.compile(r"(含义|意思|用法|用途|作用).{0,6}(不明确|不清楚|不太明白|不明白|不理解|不了解)")

COURSE_SIGNAL_WORDS = [
    "for", "while", "do", "if", "else", "switch", "case", "break", "continue",
    "scanf", "printf", "gets", "puts", "fopen", "fclose", "fprintf", "fscanf",
    "%d", "%f", "%c", "%s", "\\t", "\\n", "\\0",
    "数组", "字符串", "字符数组", "指针", "函数", "递归", "循环", "分支", "进制",
    "变量", "常量", "数据类型", "结构体", "链表", "文件", "输入", "输出",
    "ASCII", "float", "double", "int", "char"
]

# 用于“疑问/问题”标题后的续行枚举条目识别
ENUM_LINE_RE = re.compile(
    r"^\s*(?:\(?\d{1,3}\)?[\.、\)]|[一二三四五六七八九十]{1,3}[、\.．\)])\s*"
)

def strip_enum_prefix(line: str) -> str:
    return re.sub(ENUM_LINE_RE, "", (line or "")).strip()

def is_marker_header_line(line: str) -> bool:
    """
    识别：单独一行的 ‘疑问/问题/困惑/疑惑’（可带冒号）
    """
    s = (line or "").strip()
    if not s:
        return False
    s2 = re.sub(r"\s*\|\s*", " ", s).strip()
    s2n = s2.replace(" ", "")
    s2n = re.sub(r"[：:]+$", "", s2n)
    return s2n in ("疑问", "问题", "困惑", "疑惑")

# 避免“心得：...”误抽取
XINDE_PREFIX_RE = re.compile(r"^\s*心得\s*[:：]")

@dataclass
class Candidate:
    source_line_no: int
    sent_no: int
    text: str
    triggers: List[str]
    in_q_section: bool
    context_prev: Optional[str]
    context_next: Optional[str]
    llm: Optional[Dict[str, Any]] = None

def is_negative_line(line: str) -> bool:
    l = line.strip().lower()
    for pat in NEGATIVE_PATTERNS:
        if re.search(pat, l, flags=re.IGNORECASE):
            return True
    return False

# def looks_like_section_title(line: str) -> bool:
#     s = line.strip()
#     if not s:
#         return False
#     # 题目行如果含问号/明显疑问词，就不要当标题
#     if "？" in s or any(w in s for w in ["如何", "怎么", "为什么", "是什么", "是否", "能否", "可否", "吗"]):
#         return False
#
#     title_like = s.replace("：", "").replace(":", "").strip()
#     if title_like in Q_SECTION_HINTS or title_like in EXCLUDE_SECTION_HINTS or title_like in END_SECTION_HINTS:
#         return True
#
#     if re.match(r"^\s*(第?\s*[一二三四五六七八九十0-9]{1,2}\s*[、\.．:：\)])\s*", s):
#         return True
#     if re.match(r"^\s*(\(?\d{1,2}\)?\s*[、\.．:：\)])\s*", s):
#         return True
#
#     return False
def looks_like_section_title(line: str, in_q_section: bool) -> bool:
    s = line.strip()
    if not s:
        return False
    # 题目行如果含问号/明显疑问词，就不要当标题
    if "？" in s or any(w in s for w in ["如何", "怎么", "为什么", "是什么", "是否", "能否", "可否", "吗"]):
        return False

    title_like = s.replace("：", "").replace(":", "").strip()
    if title_like in Q_SECTION_HINTS or title_like in EXCLUDE_SECTION_HINTS or title_like in END_SECTION_HINTS:
        return True

    if re.match(r"^\s*(第?\s*[一二三四五六七八九十0-9]{1,2}\s*[、\.．:：\)])\s*", s):
        return True
    if re.match(r"^\s*(\(?\d{1,2}\)?\s*[、\.．:：\)])\s*", s):
        # 在“疑问段”里，1./2. 更可能是问题列表项，不当标题
        return False if in_q_section else True

    return False


def update_section_flag(line: str, in_q_section: bool) -> bool:
    s = line.strip()
    if any(h in s for h in END_SECTION_HINTS):
        return False
    if any(h in s for h in EXCLUDE_SECTION_HINTS):
        return False
    if any(h in s for h in Q_SECTION_HINTS):
        return True
    return in_q_section

def split_into_sentences_keep_q(line: str) -> List[str]:
    if not line:
        return []
    s = re.sub(r"[ \t]+", " ", line).strip()
    if not s:
        return []
    s = s.replace("?", "？")

    parts: List[str] = []
    start = 0
    for m in re.finditer(r"[。；;！!？]\s*", s):
        end = m.end()
        seg = s[start:end].strip()
        if seg:
            parts.append(seg)
        start = end
    last = s[start:].strip()
    if last:
        parts.append(last)

    # 再做轻度编号切分
    final_parts: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        sub = re.split(r"(?:\s+|^)(?=(?:\(?\d{1,2}\)?[\.、\)])\s*)", p)
        sub = [x.strip() for x in sub if x.strip()]
        if len(sub) <= 1:
            final_parts.append(p)
        else:
            buf = ""
            for x in sub:
                if re.fullmatch(r"\(?\d{1,2}\)?[\.、\)]", x):
                    buf = x
                else:
                    if buf:
                        final_parts.append((buf + " " + x).strip())
                        buf = ""
                    else:
                        final_parts.append(x)
            if buf:
                final_parts.append(buf)

    return [x for x in final_parts if len(x) >= 2]

ENUM_MARK_RE = re.compile(r"(?:^|\s)(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*")

def split_marker_tail(tail: str) -> List[str]:
    tail = tail.strip()
    if not tail:
        return []
    tail = tail.replace("?", "？")

    parts = split_into_sentences_keep_q(tail)
    if not parts:
        parts = [tail]

    refined: List[str] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue

        idxs = [m.start() for m in ENUM_MARK_RE.finditer(p)]
        if len(idxs) <= 1:
            refined.append(p)
            continue

        idxs.append(len(p))
        for a, b in zip(idxs, idxs[1:]):
            seg = p[a:b].strip()
            seg = re.sub(r"^\s*(\(?\d{1,2}\)?[\.、\)]|\(?[一二三四五六七八九十]{1,3}\)?[、\.．\)])\s*", "", seg)
            seg = seg.strip()
            if seg:
                refined.append(seg)

    out = []
    for x in refined:
        x = re.sub(r"[ \t]+", " ", x).strip()
        # 如果明显是问句但没问号，补一个（仅对强词）
        if "？" not in x and any(w in x for w in ["什么", "怎么", "如何", "为什么", "是否", "能否", "可否"]):
            x = x + "？"
        if len(x) >= 2:
            out.append(x)
    return out

def extract_marker_contents(text: str) -> List[str]:
    t = re.sub(r"\s*\|\s*", " | ", text)
    m = re.search(
        rf"({'|'.join(map(re.escape, MARKER_WORDS))})\s*(?:[:：]\s*|\|\s*[:：]?\s*\|\s*)(.+)$",
        t
    )
    if not m:
        return []
    tail = m.group(2).strip()
    tail = re.sub(r"^[\|\s:：]+", "", tail)
    tail = re.sub(r"\s*\|\s*", " ", tail).strip()
    if not tail:
        return []
    return split_marker_tail(tail)

def score_triggers(sent: str, in_q_section: bool) -> Tuple[List[str], int]:
    triggers: List[str] = []
    score = 0
    sent2 = sent.replace("?", "？")

    if "？" in sent2:
        triggers.append("HAS_QUESTION_MARK")
        score += 3

    for w in QUESTION_WORDS:
        if w in sent2:
            triggers.append(f"QWORD:{w}")
            score += 2
            break

    # v2.3：提高困惑词权重
    for w in CONFUSION_WORDS:
        if w in sent2:
            triggers.append(f"CONFUSE:{w}")
            score += 2 if not in_q_section else 3
            break

    if any(h in sent2 for h in WEAK_QUESTION_HINTS):
        triggers.append("WEAK_HINT")
        score += 1 if in_q_section else 0

    if in_q_section and score > 0:
        score += 1
        triggers.append("IN_Q_SECTION")

    return triggers, score

def extract_candidates_v2_3(lines: List[str]) -> List[Candidate]:
    cands: List[Candidate] = []
    in_q_section = False

    # 续行枚举模式：遇到“问题/疑问/困惑/疑惑”标题 或 “问题：xxx”显式marker后可开启
    marker_cont_active = False

    for i, raw_line in enumerate(lines):
        line = (raw_line or "").strip()
        if not line:
            marker_cont_active = False
            continue
        if is_negative_line(line):
            marker_cont_active = False
            continue

        prev_line = lines[i - 1].strip() if i - 1 >= 0 else None
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None

        # 标题更新 section
        # if looks_like_section_title(line):
        if looks_like_section_title(line, in_q_section):
            in_q_section = update_section_flag(line, in_q_section)
            marker_cont_active = False
            continue

        # 只在疑问段抽取（硬门控）
        if not in_q_section:
            marker_cont_active = False
            continue

        # -------------------------
        # C) 避免“心得：...区别...”误抽取
        # 心得行：若无问号/困惑词/显式marker，就跳过整行
        # -------------------------
        if XINDE_PREFIX_RE.search(line):
            tmp_triggers, _ = score_triggers(line, in_q_section=True)
            has_explicit_marker = bool(extract_marker_contents(line))
            has_confuse = any(t.startswith("CONFUSE:") for t in tmp_triggers)
            has_q = ("HAS_QUESTION_MARK" in tmp_triggers) or any(t.startswith("QWORD:") for t in tmp_triggers)
            if (not has_explicit_marker) and (not has_confuse) and ("HAS_QUESTION_MARK" not in tmp_triggers):
                # 纯心得总结句：不抽（即使包含“区别/作用”等词）
                continue

        # 元数据行过滤（但 marker / 问号保留）
        if any(f in line for f in META_FIELD_HINTS):
            has_qmark = ("?" in line or "？" in line)
            has_marker = bool(extract_marker_contents(line))
            if not has_qmark and not has_marker:
                continue

        # -------------------------
        # B1) “疑问/问题/困惑/疑惑” 单独一行标题 -> 开启续行吸收
        # -------------------------
        if is_marker_header_line(line):
            marker_cont_active = True
            continue

        # -------------------------
        # B2) 显式 marker（同一行：问题：xxx）
        # -------------------------
        marker_parts = extract_marker_contents(line)
        if marker_parts:
            for j, p in enumerate(marker_parts):
                p = p.strip()
                if not p:
                    continue
                triggers, _ = score_triggers(p, in_q_section=True)
                triggers = (triggers or []) + ["MARKER:EXPLICIT", "IN_Q_SECTION"]
                cands.append(Candidate(
                    source_line_no=i,
                    sent_no=j,
                    text=p.replace("?", "？"),
                    triggers=sorted(list(set(triggers))),
                    in_q_section=True,
                    context_prev=prev_line,
                    context_next=next_line,
                    llm=None
                ))
            # 显式 marker 后面很可能继续 2. / 3. ...
            marker_cont_active = True
            continue

        # -------------------------
        # B3) marker 续行枚举条目（允许名词化）
        # -------------------------
        if marker_cont_active and ENUM_LINE_RE.search(line):
            body = strip_enum_prefix(line).replace("?", "？").strip()
            if body and len(body) >= 2:
                triggers, _ = score_triggers(body, in_q_section=True)
                triggers = (triggers or []) + ["MARKER:CONT", "IN_Q_SECTION"]

                # 续行条目允许弱/名词化，但仍做底线：太短不要
                if len(body) >= 3:
                    cands.append(Candidate(
                        source_line_no=i,
                        sent_no=0,
                        text=body,
                        triggers=sorted(list(set(triggers))),
                        in_q_section=True,
                        context_prev=prev_line,
                        context_next=next_line,
                        llm=None
                    ))
            continue

        # marker 续行遇到非枚举行：若该行无任何疑问信号，则关闭（避免吸入心得）
        if marker_cont_active:
            tmp_triggers, _ = score_triggers(line.replace("?", "？"), in_q_section=True)
            has_signal = (
                ("HAS_QUESTION_MARK" in tmp_triggers)
                or any(t.startswith("QWORD:") for t in tmp_triggers)
                or any(t.startswith("CONFUSE:") for t in tmp_triggers)
            )
            if not has_signal:
                marker_cont_active = False
            # 不 continue，让下面正常分句逻辑处理本行（如果本行就是问句，也能抽到）

        # -------------------------
        # 原有：分句抽取
        # -------------------------
        sents = split_into_sentences_keep_q(line)
        for j, sent in enumerate(sents):
            sent = sent.strip()
            if not sent:
                continue
            sent = sent.replace("?", "？")

            if WEAK_CONFUSE_PAT.search(sent) and "？" not in sent and not any(w in sent for w in QUESTION_WORDS):
                continue

            if len(sent) <= 4 and ("？" not in sent):
                continue

            triggers, score = score_triggers(sent, in_q_section=True)
            if score < 2:
                continue

            only_weak = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                (any(t.startswith("CONFUSE:") for t in triggers) or "WEAK_HINT" in triggers)
            )

            confuse_only = (
                ("HAS_QUESTION_MARK" not in triggers) and
                (not any(t.startswith("QWORD:") for t in triggers)) and
                any(t.startswith("CONFUSE:") for t in triggers)
            )

            sent_l = sent.lower()
            has_course_signal = any(k.lower() in sent_l for k in COURSE_SIGNAL_WORDS)

            if only_weak:
                if in_q_section and confuse_only and has_course_signal and len(sent) >= 8:
                    triggers.append("ALLOW:CONFUSE_WITH_SIGNAL")
                else:
                    continue

            cands.append(Candidate(
                source_line_no=i,
                sent_no=j,
                text=sent,
                triggers=sorted(list(set(triggers))),
                in_q_section=True,
                context_prev=prev_line,
                context_next=next_line,
                llm=None
            ))

    # 去重
    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        if c.text not in seen:
            uniq.append(c)
            seen.add(c.text)
    return uniq

# =====================
# 3) strict LLM 兜底抽取（整篇喂给LLM，但允许输出[]）
# =====================

def _safe_json_loads(s: str):
    s = (s or "").strip()
    try:
        return json.loads(s)
    except Exception:
        a = s.find("[")
        b = s.rfind("]")
        if a != -1 and b != -1 and b > a:
            return json.loads(s[a:b + 1])
        raise

def call_llm_extract_strict(full_text: str, file_name: str, timeout: int = 120) -> List[Dict[str, Any]]:
    api_key = OPENAI_API_KEY
    base_url = OPENAI_BASE_URL
    model = OPENAI_MODEL

    if not api_key or api_key == "sk-REPLACE_ME":
        raise RuntimeError("Missing OPENAI_API_KEY. Please set env OPENAI_API_KEY.")

    import requests
    url = base_url.rstrip("/") + "/v1/chat/completions"

    system = (
        "你是信息抽取助手。任务：从学生自学报告全文中抽取“可用于RAG问答”的疑问/困惑点。\n"
        "【只允许抽取两类】\n"
        "A) QUESTION：明确问句，满足其一即可：\n"
        "   - 含问号（? 或 ？）\n"
        "   - 或包含疑问词：什么/怎么/如何/为什么/是否/能否/可否/区别/联系/原因/作用/意义/什么时候/哪里/多少/吗\n"
        "B) CONFUSION：明确困惑陈述，必须同时满足：\n"
        "   - 含困惑表达：不懂/不理解/没理解/不太懂/不太明白/不清楚/疑惑/困惑/搞不懂/不知道/不确定/不熟悉/把握不好/掌握不好/卡壳\n"
        "   - 且指向具体对象（概念/语法/函数/例子/知识点），不是空泛的“学得不好/有困难”。\n"
        "【严格排除】\n"
        "1) 只列出主题/知识点列表但没有问法或困惑动词，一律不要抽取。\n"
        "2) 纯心得/收获/总结/计划，不要抽取。\n"
        "3) 过于笼统无法形成问答的句子（如“编写代码有困难”）默认不抽取，除非后面紧跟具体困难点。\n"
        "【输出要求】\n"
        "- 只输出严格JSON数组（不要Markdown/不要解释）。\n"
        "- evidence 必须从原文复制（<=60字优先），用于定位。\n"
        "- question 尽量贴近原文，不要做规范化扩写。\n"
        "- 若没有符合条件的内容，输出 []。\n"
    )

    user = f"""文件名：{file_name}

【文档全文（按行拼接）】
{full_text}

【输出JSON数组格式示例】
[
  {{
    "question": "对二进制、十进制、八进制、十六进制的换算不懂",
    "evidence": "对二进制、十进制、八进制、十六进制的换算不懂",
    "category": "CONFUSION",
    "confidence": 0.9,
    "rationale": "包含“不懂”且指向具体知识点“进制换算”"
  }},
  {{
    "question": "printf括号中的“%d”作用是什么？",
    "evidence": "printf括号中的“%d”作用是什么",
    "category": "QUESTION",
    "confidence": 0.95,
    "rationale": "明确问句，包含“作用是什么”"
  }}
]
"""

    payload = {
        "model": model,
        "temperature": 0.0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    }

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if r.status_code >= 400:
        raise RuntimeError(f"LLM HTTP {r.status_code}: {r.text[:800]}")

    data = r.json()
    content = data["choices"][0]["message"]["content"]
    extracted = _safe_json_loads(content)
    if not isinstance(extracted, list):
        raise ValueError("LLM output is not a JSON array.")
    return extracted

def _locate_in_lines(lines: List[str], evidence: str) -> int:
    if not evidence:
        return -1
    ev = re.sub(r"[ \t]+", " ", evidence).strip()
    if len(ev) < 2:
        return -1
    for i, ln in enumerate(lines):
        if ev in ln:
            return i
    head = ev[:12]
    if len(head) >= 4:
        for i, ln in enumerate(lines):
            if head in ln:
                return i
    return -1

def build_llm_candidates(lines: List[str], llm_items: List[Dict[str, Any]]) -> List[Candidate]:
    cands: List[Candidate] = []
    for it in llm_items:
        q = (it.get("question") or "").strip()
        if not q:
            continue

        evidence = (it.get("evidence") or q).strip()
        line_no = _locate_in_lines(lines, evidence)

        prev_line = lines[line_no - 1].strip() if (line_no >= 1) else None
        next_line = lines[line_no + 1].strip() if (line_no >= 0 and line_no + 1 < len(lines)) else None

        cands.append(Candidate(
            source_line_no=line_no,
            sent_no=0,
            text=q.replace("?", "？"),
            triggers=["LLM_FALLBACK_STRICT"],
            in_q_section=False,
            context_prev=prev_line,
            context_next=next_line,
            llm={
                "category": it.get("category", ""),
                "rationale": it.get("rationale", ""),
                "evidence": evidence,
                "confidence": it.get("confidence", None),
            }
        ))

    uniq: List[Candidate] = []
    seen = set()
    for c in cands:
        if c.text not in seen:
            uniq.append(c)
            seen.add(c.text)
    return uniq

# =====================
# 4) 主流程：规则优先 + 0候选兜底LLM
# =====================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="file or folder containing .docx/.pdf")
    parser.add_argument("--out", default="question_candidates_v2_3_fallback_fixed.json", help="output json filename")
    parser.add_argument("--max_chars", type=int, default=24000, help="truncate text for LLM prompt")
    parser.add_argument("--llm_timeout", type=int, default=OPENAI_TIMEOUT, help="LLM request timeout seconds")
    parser.add_argument("--max_preview", type=int, default=0, help="print top N candidates per file for quick check")
    parser.add_argument("--debug_file", default="", help="只调试打印包含该关键词的文件名（如 符丽）")
    parser.add_argument("--debug_kw", default="问题", help="调试时定位关键词（默认：问题）")
    args = parser.parse_args()

    out_path = make_nonconflicting_path(args.out)

    paths = iter_files(args.path)
    if not paths:
        raise FileNotFoundError(f"No .docx/.pdf found under: {args.path}")

    results = []
    for p in paths:
        try:
            lines = read_lines(p)
            # ===== DEBUG PRINT =====
            if args.debug_file and (args.debug_file in os.path.basename(p)):
                print("\n" + "=" * 80)
                print("[DEBUG] file:", p)
                print("[DEBUG] total lines:", len(lines))

                # 先把前 50 行都列出来（很多时候问题就在这里）
                print("\n[DEBUG] first 50 lines (repr):")
                for idx, ln in enumerate(lines[:50]):
                    print(f"{idx:03d}: {repr(ln)}")

                # 再找包含关键词的位置（默认找“问题”）
                hits = [i for i, ln in enumerate(lines) if args.debug_kw in ln]
                print(f"\n[DEBUG] hit indices for kw='{args.debug_kw}':", hits)

                # 打印命中附近上下文
                for hi in hits[:10]:
                    print("\n[DEBUG] around index", hi)
                    for j in range(max(0, hi - 5), min(len(lines), hi + 10)):
                        print(f"{j:03d}: {repr(lines[j])}")
                print("=" * 80 + "\n")
            # ===== DEBUG PRINT END =====
            # 1) 规则抽取
            rule_cands = extract_candidates_v2_3(lines)

            # 2) 若规则为0，调用LLM兜底（LLM允许输出[]）
            final_cands = rule_cands
            used_fallback = False

            if len(rule_cands) == 0:
                joined = "\n".join(lines)
                if len(joined) > args.max_chars:
                    joined = joined[:args.max_chars] + "\n(后续内容因长度限制已截断)"
                llm_items = call_llm_extract_strict(joined, os.path.basename(p), timeout=args.llm_timeout)
                llm_cands = build_llm_candidates(lines, llm_items)
                final_cands = llm_cands
                used_fallback = True

            results.append({
                "file": os.path.abspath(p),
                "num_lines": len(lines),
                "num_candidates": len(final_cands),
                "candidates": [asdict(x) for x in final_cands],
                "fallback_used": used_fallback,
            })

            tag = "RULE" if not used_fallback else "FALLBACK_LLM"
            print(f"[OK] {os.path.basename(p)} lines={len(lines)} candidates={len(final_cands)} mode={tag}")

            if args.max_preview > 0:
                for x in final_cands[:args.max_preview]:
                    extra = ""
                    if x.llm:
                        extra = f"  [cat={x.llm.get('category')}, conf={x.llm.get('confidence')}]"
                    print(f"  - (L{x.source_line_no}:{x.sent_no}) {x.text} triggers={x.triggers}{extra}")

        except Exception as e:
            print(f"[FAIL] {p}: {e}")
            results.append({"file": os.path.abspath(p), "error": str(e)})

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nSaved: {out_path} ({len(results)} files)")

if __name__ == "__main__":
    main()
